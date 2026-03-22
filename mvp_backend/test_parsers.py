from __future__ import annotations

import shutil
import sys
import types
import unittest
import uuid
from pathlib import Path
from unittest.mock import patch

try:
    from docx import Document
except ImportError:  # pragma: no cover - optional dependency in local env
    Document = None
    docx_stub = types.ModuleType('docx')

    class _MissingDocument:
        def __init__(self, *args, **kwargs):
            raise RuntimeError('python-docx is not installed in the current environment')

    docx_stub.Document = _MissingDocument
    sys.modules['docx'] = docx_stub

try:
    import pdfplumber  # type: ignore  # noqa: F401
except ImportError:  # pragma: no cover - optional dependency in local env
    pdfplumber_stub = types.ModuleType('pdfplumber')

    class _MissingPdfContext:
        def __enter__(self):
            raise RuntimeError('pdfplumber is not installed in the current environment')

        def __exit__(self, exc_type, exc, tb):
            return None

    def _missing_pdf_open(*args, **kwargs):
        return _MissingPdfContext()

    pdfplumber_stub.open = _missing_pdf_open
    sys.modules['pdfplumber'] = pdfplumber_stub

try:
    import openpyxl  # type: ignore  # noqa: F401
except ImportError:  # pragma: no cover - optional dependency in local env
    openpyxl = None

BACKEND_DIR = Path(__file__).resolve().parent
if str(BACKEND_DIR) not in sys.path:
    sys.path.insert(0, str(BACKEND_DIR))

try:
    import pydantic  # type: ignore  # noqa: F401
except ModuleNotFoundError:
    pydantic_stub = types.ModuleType('pydantic')

    class _FieldInfo:
        def __init__(self, default=None, default_factory=None):
            self.default = default
            self.default_factory = default_factory

    def Field(default=None, default_factory=None):
        return _FieldInfo(default=default, default_factory=default_factory)

    class BaseModel:
        def __init__(self, **data):
            annotations = {}
            for cls in reversed(self.__class__.mro()):
                annotations.update(getattr(cls, '__annotations__', {}))
            for key in annotations:
                if key in data:
                    value = data[key]
                else:
                    class_value = getattr(self.__class__, key, None)
                    if isinstance(class_value, _FieldInfo):
                        if class_value.default_factory is not None:
                            value = class_value.default_factory()
                        else:
                            value = class_value.default
                    else:
                        value = class_value
                setattr(self, key, value)

        def dict(self):
            return dict(self.__dict__)

        def model_dump(self):
            return dict(self.__dict__)

        def copy(self, update=None):
            payload = dict(self.__dict__)
            payload.update(update or {})
            return self.__class__(**payload)

        def model_copy(self, update=None):
            return self.copy(update=update)

    pydantic_stub.BaseModel = BaseModel
    pydantic_stub.Field = Field
    sys.modules['pydantic'] = pydantic_stub

try:
    import fastapi  # type: ignore  # noqa: F401
except ModuleNotFoundError:
    fastapi_stub = types.ModuleType('fastapi')
    fastapi_security_stub = types.ModuleType('fastapi.security')

    class APIRouter:
        def post(self, *args, **kwargs):
            def decorator(func):
                return func
            return decorator

        def get(self, *args, **kwargs):
            def decorator(func):
                return func
            return decorator

        def delete(self, *args, **kwargs):
            def decorator(func):
                return func
            return decorator

    class HTTPException(Exception):
        def __init__(self, status_code: int, detail: str, headers=None):
            self.status_code = status_code
            self.detail = detail
            self.headers = headers or {}
            super().__init__(detail)

    class UploadFile:
        def __init__(self, *args, **kwargs):
            self.filename = kwargs.get('filename')

    class HTTPAuthorizationCredentials:
        def __init__(self, credentials: str = ''):
            self.credentials = credentials

    class HTTPBearer:
        def __init__(self, *args, **kwargs):
            self.auto_error = kwargs.get('auto_error', True)

        def __call__(self, *args, **kwargs):
            return None

    def Depends(value=None):
        return value

    def File(*args, **kwargs):
        return None

    def Form(*args, **kwargs):
        return None

    fastapi_stub.APIRouter = APIRouter
    fastapi_stub.Depends = Depends
    fastapi_stub.File = File
    fastapi_stub.Form = Form
    fastapi_stub.HTTPException = HTTPException
    fastapi_stub.UploadFile = UploadFile
    fastapi_stub.status = types.SimpleNamespace(HTTP_401_UNAUTHORIZED=401)
    fastapi_security_stub.HTTPAuthorizationCredentials = HTTPAuthorizationCredentials
    fastapi_security_stub.HTTPBearer = HTTPBearer
    sys.modules['fastapi'] = fastapi_stub
    sys.modules['fastapi.security'] = fastapi_security_stub

from models import ParsedFile, ParsedSheet, RepairApplyPayload, RepairPreviewPayload, TargetField
from parsers import ParseError, _build_form_model, _resolve_generic_form_layout_source, parse_file, parse_target_schema, resolve_generation_source
from pdf_zoning import classify_pdf_document_zones
from document_parser import _merge_ocr_image_results, _suppress_consumed_group_fragments
from form_layout import (
    _extract_scalar_from_table_row,
    _extract_scalars_from_layout_lines,
    _extract_table_rows_as_form,
    _parse_kv_line,
    _resolve_fatca_group,
    understand_generic_form,
)
from routes import _build_form_explainability, repair_apply, repair_preview


class FakeWorksheet:
    def __init__(self, rows: list[tuple[object, ...]]) -> None:
        self._rows = rows

    def iter_rows(self, values_only: bool = True):
        return iter(self._rows)


class FakeWorkbook:
    def __init__(self, sheets: dict[str, FakeWorksheet]) -> None:
        self._sheets = sheets
        self.sheetnames = list(sheets.keys())

    def __getitem__(self, sheet_name: str) -> FakeWorksheet:
        return self._sheets[sheet_name]

    def close(self) -> None:
        return None


class FakePdfPage:
    def __init__(self, *, text: str, words: list[dict[str, object]] | None = None, tables: list[list[list[str]]] | None = None) -> None:
        self._text = text
        self._words = words or []
        self._tables = tables or []

    def extract_text(self):
        return self._text

    def extract_words(self):
        return list(self._words)

    def extract_tables(self):
        return list(self._tables)


class FakePdfContext:
    def __init__(self, pages: list[FakePdfPage]) -> None:
        self.pages = pages

    def __enter__(self):
        return self

    def __exit__(self, exc_type, exc, tb):
        return None


class FakeLlamaDocument:
    def __init__(self, text: str, metadata: dict[str, object] | None = None) -> None:
        self.text = text
        self.metadata = metadata or {}


@unittest.skipIf(Document is None, 'python-docx is not installed in the current environment')
class DocumentParserTests(unittest.TestCase):
    def setUp(self) -> None:
        self.test_root = BACKEND_DIR / '.test_runtime' / 'documents' / str(uuid.uuid4())
        self.test_root.mkdir(parents=True, exist_ok=True)

    def tearDown(self) -> None:
        shutil.rmtree(self.test_root, ignore_errors=True)

    def test_parse_target_schema_allows_empty_array_alongside_string_arrays(self) -> None:
        target_json = """
        {
          "input": [
            {
              "organizationName": "ООО \\"Рога и копыта\\"",
              "innOrKio": "1234567890",
              "isResidentRF": "NOWHERE",
              "isTaxResidencyOnlyRF": "NO",
              "fatcaBeneficiaryOptionList": [
                "IS_DISREGARDED_ENTITY"
              ]
            },
            {
              "organizationName": "ООО \\"Иноагент\\"",
              "innOrKio": "0987654321",
              "isResidentRF": "YES",
              "isTaxResidencyOnlyRF": "NO",
              "fatcaBeneficiaryOptionList": [
                "IS_DISREGARDED_ENTITY",
                "IS_FATCA_FOREIGN_INSTITUTE"
              ]
            },
            {
              "organizationName": "ООО \\"Наши люди\\"",
              "innOrKio": "6789054321",
              "isResidentRF": "YES",
              "isTaxResidencyOnlyRF": "YES",
              "fatcaBeneficiaryOptionList": []
            }
          ]
        }
        """

        target_fields, payload, schema, summary = parse_target_schema(target_json)

        self.assertTrue(target_fields)
        self.assertIsInstance(payload, list)
        self.assertEqual(schema["type"], "array")
        self.assertEqual(summary["root_type"], "array")

    def test_parse_target_schema_unwraps_single_input_array_wrapper(self) -> None:
        target_json = """
        {
          "input": [
            {
              "organizationName": "ООО \\"Рога и копыта\\"",
              "innOrKio": "1234567890",
              "isResidentRF": "NOWHERE"
            }
          ]
        }
        """

        target_fields, payload, schema, summary = parse_target_schema(target_json)

        self.assertIsInstance(payload, list)
        self.assertEqual(schema["type"], "array")
        self.assertTrue(summary["root_is_array"])
        self.assertEqual([field.name for field in target_fields], ["organizationName", "innOrKio", "isResidentRF"])

    def test_docx_table_is_parsed_without_generic_no_columns_warning(self) -> None:
        path = self.test_root / 'table.docx'
        doc = Document()
        table = doc.add_table(rows=2, cols=2)
        table.rows[0].cells[0].text = 'customerName'
        table.rows[0].cells[1].text = 'amount'
        table.rows[1].cells[0].text = 'Alice'
        table.rows[1].cells[1].text = '10'
        doc.save(path)

        parsed = parse_file(path, path.name)

        self.assertEqual(parsed.columns, ['customerName', 'amount'])
        self.assertEqual(parsed.rows, [{'customerName': 'Alice', 'amount': '10'}])
        self.assertNotIn('No columns detected in the file.', parsed.warnings)

    def test_docx_text_fallback_returns_document_warning_without_generic_no_columns_warning(self) -> None:
        path = self.test_root / 'text.docx'
        doc = Document()
        doc.add_paragraph('Just plain text without tables')
        doc.save(path)

        parsed = parse_file(path, path.name)

        self.assertEqual(parsed.columns, [])
        self.assertTrue(any('No tables found in DOCX' in warning for warning in parsed.warnings))
        self.assertEqual(parsed.content_type, 'text')
        self.assertEqual(parsed.extraction_status, 'text_extracted')
        self.assertNotIn('No columns detected in the file.', parsed.warnings)

    def test_docx_paragraph_kv_rows_are_reconstructed_into_wide_table(self) -> None:
        path = self.test_root / 'paragraph_format.docx'
        doc = Document()
        doc.add_paragraph('Paragraph Format')
        doc.add_paragraph(
            'dealId: DEAL_12345, dealName: Проект Alpha, creationDate: 2025-07-15, '
            'stage: Переговоры, revenue: 1000000, organization: ООО Организация, '
            'product: Продукт A, responsible: Иванов Иван'
        )
        doc.add_paragraph(
            'dealId: DEAL_67890, dealName: Проект Beta, creationDate: 2025-06-10, '
            'stage: Закрыта успешно, revenue: 750000, organization: ООО Техно, '
            'product: Сервис Y, responsible: Сидоров Алексей'
        )
        doc.save(path)

        parsed = parse_file(path, path.name)

        self.assertEqual(parsed.content_type, 'table')
        self.assertEqual(parsed.document_mode, 'data_table_mode')
        self.assertEqual(
            parsed.columns,
            ['dealId', 'dealName', 'creationDate', 'stage', 'revenue', 'organization', 'product', 'responsible'],
        )
        self.assertEqual(parsed.rows[0]['dealId'], 'DEAL_12345')
        self.assertEqual(parsed.rows[0]['organization'], 'ООО Организация')
        self.assertEqual(parsed.rows[1]['responsible'], 'Сидоров Алексей')
        self.assertTrue(any('Reconstructed tabular preview' in warning for warning in parsed.warnings))

    def test_docx_single_line_pipe_payload_is_reconstructed_into_repeated_rows(self) -> None:
        path = self.test_root / 'one_line_format.docx'
        doc = Document()
        doc.add_paragraph('One Line Format')
        doc.add_paragraph(
            'DEAL_12345 | Проект Alpha | 2025-07-15 | Переговоры | 1000000 | ООО Организация | '
            'Продукт A | Иванов Иван | DEAL_67890 | Проект Beta | 2025-06-10 | '
            'Закрыта успешно | 750000 | ООО Техно | Сервис Y | Сидоров Алексей'
        )
        doc.save(path)

        parsed = parse_file(path, path.name)

        self.assertEqual(parsed.content_type, 'table')
        self.assertEqual(parsed.document_mode, 'data_table_mode')
        self.assertEqual(parsed.columns, ['identifier', 'name', 'date', 'stage', 'amount', 'organization', 'product', 'responsible'])
        self.assertEqual(len(parsed.rows), 2)
        self.assertEqual(parsed.rows[0]['identifier'], 'DEAL_12345')
        self.assertEqual(parsed.rows[0]['amount'], '1000000')
        self.assertEqual(parsed.rows[1]['organization'], 'ООО Техно')
        self.assertEqual(parsed.rows[1]['responsible'], 'Сидоров Алексей')
        self.assertTrue(any('Reconstructed tabular preview' in warning for warning in parsed.warnings))

    def test_docx_with_multiple_regular_tables_stays_tabular(self) -> None:
        path = self.test_root / 'three_tables.docx'
        doc = Document()

        inventory = doc.add_table(rows=3, cols=3)
        inventory.rows[0].cells[0].text = 'SKU'
        inventory.rows[0].cells[1].text = 'Item Name'
        inventory.rows[0].cells[2].text = 'Price'
        inventory.rows[1].cells[0].text = 'SKU001'
        inventory.rows[1].cells[1].text = 'Laptop'
        inventory.rows[1].cells[2].text = '45000'
        inventory.rows[2].cells[0].text = 'SKU002'
        inventory.rows[2].cells[1].text = 'Mouse'
        inventory.rows[2].cells[2].text = '900'

        warehouse = doc.add_table(rows=3, cols=2)
        warehouse.rows[0].cells[0].text = 'Warehouse'
        warehouse.rows[0].cells[1].text = 'Stock'
        warehouse.rows[1].cells[0].text = 'WH-A'
        warehouse.rows[1].cells[1].text = '18'
        warehouse.rows[2].cells[0].text = 'WH-B'
        warehouse.rows[2].cells[1].text = '24'

        employees = doc.add_table(rows=3, cols=2)
        employees.rows[0].cells[0].text = 'Employee'
        employees.rows[0].cells[1].text = 'Department'
        employees.rows[1].cells[0].text = 'Ann'
        employees.rows[1].cells[1].text = 'Sales'
        employees.rows[2].cells[0].text = 'Ben'
        employees.rows[2].cells[1].text = 'Ops'

        doc.save(path)

        parsed = parse_file(path, path.name)

        self.assertEqual(parsed.content_type, 'table')
        self.assertEqual(parsed.document_mode, 'data_table_mode')
        self.assertEqual(len(parsed.sheets), 3)
        self.assertTrue(any('Found 3 tables in DOCX.' in warning for warning in parsed.warnings))
        self.assertFalse(any('form-like layout document' in warning for warning in parsed.warnings))

    def test_docx_form_text_extracts_kv_pairs_and_candidates(self) -> None:
        path = self.test_root / 'form.docx'
        doc = Document()
        doc.add_paragraph('ФИО: Иванов Иван')
        doc.add_paragraph('Дата рождения: 01.01.1990')
        doc.add_paragraph('Страна налогового резидентства: Германия')
        doc.save(path)

        parsed = parse_file(path, path.name)
        columns, rows, warnings = resolve_generation_source(parsed)

        self.assertEqual(parsed.content_type, 'form')
        self.assertEqual(parsed.extraction_status, 'text_extracted')
        self.assertEqual([pair.label for pair in parsed.kv_pairs], ['ФИО', 'Дата рождения', 'Страна налогового резидентства'])
        self.assertEqual(columns, ['ФИО', 'Дата рождения', 'Страна налогового резидентства'])
        self.assertEqual(rows, [{'ФИО': 'Иванов Иван', 'Дата рождения': '01.01.1990', 'Страна налогового резидентства': 'Германия'}])
        self.assertIn('Generated mapping from extracted fields/text candidates.', warnings)

    def test_docx_single_inline_profile_extracts_multiple_kv_pairs(self) -> None:
        path = self.test_root / 'single_inline_profile.docx'
        doc = Document()
        doc.add_paragraph(
            'Это тестовый документ в одну строку: '
            'ФИО Иван Иванов, дата рождения 01.01.1990, телефон +7 999 123-45-67, '
            'email ivan.ivanov@example.com, адрес г. Москва, ул. Тестовая, д. 10, '
            'цель документа проверка обработки текста без переносов строк.'
        )
        doc.save(path)

        parsed = parse_file(path, path.name)
        columns, rows, warnings = resolve_generation_source(parsed)

        self.assertEqual(parsed.content_type, 'form')
        self.assertEqual(parsed.document_mode, 'form_layout_mode')
        self.assertEqual(
            [pair.label for pair in parsed.kv_pairs],
            ['ФИО', 'дата рождения', 'телефон', 'email', 'адрес', 'цель документа'],
        )
        self.assertEqual(
            columns,
            ['ФИО', 'дата рождения', 'телефон', 'email', 'адрес', 'цель документа'],
        )
        self.assertEqual(rows[0]['ФИО'], 'Иван Иванов')
        self.assertEqual(rows[0]['дата рождения'], '01.01.1990')
        self.assertEqual(rows[0]['телефон'], '+7 999 123-45-67')
        self.assertEqual(rows[0]['email'], 'ivan.ivanov@example.com')
        self.assertEqual(rows[0]['адрес'], 'г. Москва, ул. Тестовая, д. 10')
        self.assertEqual(rows[0]['цель документа'], 'проверка обработки текста без переносов строк')
        self.assertIn('Generated mapping from extracted fields/text candidates.', warnings)

    def test_docx_single_inline_profile_generate_uses_extracted_field_candidates(self) -> None:
        path = self.test_root / 'single_inline_profile_generate.docx'
        doc = Document()
        doc.add_paragraph(
            'Это тестовый документ в одну строку: '
            'ФИО Иван Иванов, дата рождения 01.01.1990, телефон +7 999 123-45-67, '
            'email ivan.ivanov@example.com, адрес г. Москва, ул. Тестовая, д. 10, '
            'цель документа проверка обработки текста без переносов строк.'
        )
        doc.save(path)

        parsed = parse_file(path, path.name)
        columns, rows, warnings = resolve_generation_source(
            parsed,
            target_fields=[
                TargetField(name='fullName', type='string'),
                TargetField(name='birthDate', type='string'),
                TargetField(name='phoneNumber', type='string'),
                TargetField(name='emailAddress', type='string'),
                TargetField(name='address', type='string'),
                TargetField(name='documentPurpose', type='string'),
            ],
        )

        self.assertEqual(parsed.document_mode, 'form_layout_mode')
        self.assertEqual(
            columns,
            ['ФИО', 'дата рождения', 'телефон', 'email', 'адрес', 'цель документа'],
        )
        self.assertEqual(rows[0]['ФИО'], 'Иван Иванов')
        self.assertTrue(any('fell back to extracted field/value candidates' in warning.lower() for warning in warnings))

    def test_txt_form_extracts_kv_pairs_and_sections(self) -> None:
        path = self.test_root / 'form.txt'
        path.write_text(
            'Анкета клиента\n\nФИО: Иванов Иван\nДата рождения: 01.01.1990\nСтрана налогового резидентства: Германия\n',
            encoding='utf-8',
        )

        parsed = parse_file(path, path.name)
        columns, rows, _warnings = resolve_generation_source(parsed)

        self.assertEqual(parsed.file_type, 'txt')
        self.assertEqual(parsed.content_type, 'form')
        self.assertEqual(parsed.extraction_status, 'text_extracted')
        self.assertGreaterEqual(len(parsed.sections), 1)
        self.assertEqual(columns, ['ФИО', 'Дата рождения', 'Страна налогового резидентства'])
        self.assertEqual(rows[0]['Страна налогового резидентства'], 'Германия')

    def test_txt_with_multiple_tables_exposes_each_table_as_sheet(self) -> None:
        path = self.test_root / 'multi_table.txt'
        path.write_text(
            'productId\tproductName\tprice\n'
            '1\tLaptop\t1000\n'
            '2\tMouse\t25\n'
            '\n'
            'orderId\tstatus\ttotal\n'
            'A-1\tpaid\t1025\n'
            'A-2\tpending\t400\n',
            encoding='utf-8',
        )

        parsed = parse_file(path, path.name)

        self.assertEqual(parsed.file_type, 'txt')
        self.assertEqual([sheet.name for sheet in parsed.sheets], ['Table 1', 'Table 2'])
        self.assertEqual(parsed.sheets[0].columns, ['productId', 'productName', 'price'])
        self.assertEqual(parsed.sheets[1].columns, ['orderId', 'status', 'total'])
        self.assertEqual(parsed.sheets[0].rows[0]['productName'], 'Laptop')
        self.assertEqual(parsed.sheets[1].rows[0]['status'], 'paid')
        self.assertTrue(any('Found 2 tables in TXT' in warning for warning in parsed.warnings))

    def test_image_like_input_reports_ocr_requirement(self) -> None:
        path = self.test_root / 'scan.png'
        path.write_bytes(b'not-a-real-png-but-extension-is-enough')

        parsed = parse_file(path, path.name)

        self.assertEqual(parsed.content_type, 'image_like')
        self.assertEqual(parsed.extraction_status, 'image_parse_not_supported_yet')
        self.assertEqual(parsed.columns, [])
        self.assertEqual(parsed.rows, [])
        self.assertTrue(any('OCR' in warning or 'image' in warning.lower() for warning in parsed.warnings))

    def test_image_like_input_can_use_external_ocr_service(self) -> None:
        path = self.test_root / 'photo.jpg'
        path.write_bytes(b'not-a-real-photo-but-extension-is-enough')

        with patch(
            'document_parser.extract_text_from_ocr_service',
            return_value={
                'text': 'Наименование организации: ООО "Рога и копыта"\nИНН/КИО: 1234567890',
                'blocks': [
                    {'id': 'ocr-line-1', 'kind': 'line', 'text': 'Наименование организации: ООО "Рога и копыта"', 'label': 'ocr'},
                    {'id': 'ocr-line-2', 'kind': 'line', 'text': 'ИНН/КИО: 1234567890', 'label': 'ocr'},
                ],
                'warnings': ['Text was extracted via the external OCR service.'],
            },
        ):
            parsed = parse_file(path, path.name)
            columns, rows, warnings = resolve_generation_source(parsed)

        self.assertIn(parsed.content_type, {'form', 'text'})
        self.assertEqual(parsed.extraction_status, 'text_extracted')
        self.assertEqual(columns, ['Наименование организации', 'ИНН/КИО'])
        self.assertEqual(rows[0]['ИНН/КИО'], '1234567890')
        self.assertTrue(any('external OCR service' in warning for warning in parsed.warnings + warnings))

    def test_docx_form_like_table_is_resolved_via_form_aware_extraction(self) -> None:
        path = self.test_root / 'form_layout.docx'
        doc = Document()
        table = doc.add_table(rows=8, cols=2)
        table.rows[0].cells[0].text = 'Наименование организации'
        table.rows[0].cells[1].text = 'ООО "Рога и копыта"'
        table.rows[1].cells[0].text = 'ИНН/КИО'
        table.rows[1].cells[1].text = '1234567890'
        table.rows[2].cells[0].text = 'Является ли выгодоприобретатель налоговым резидентом только в РФ'
        table.rows[3].cells[0].text = ' '
        table.rows[3].cells[1].text = 'ДА'
        table.rows[4].cells[0].text = 'X'
        table.rows[4].cells[1].text = 'Не являюсь налоговым резидентом ни в одном государстве'
        table.rows[5].cells[0].text = ' '
        table.rows[5].cells[1].text = 'НЕТ, является налоговым резидентом в иностранном государстве'
        table.rows[6].cells[0].text = 'FATCA статус выгодоприобретателя'
        table.rows[7].cells[0].text = 'X'
        table.rows[7].cells[1].text = 'Иностранным финансовым институтом'
        doc.save(path)

        parsed = parse_file(path, path.name)
        columns, rows, warnings = resolve_generation_source(
            parsed,
            target_fields=[
                TargetField(name='organizationName', type='string'),
                TargetField(name='innOrKio', type='string'),
                TargetField(name='isResidentRF', type='string'),
                TargetField(name='isTaxResidencyOnlyRF', type='string'),
                TargetField(name='fatcaBeneficiaryOptionList', type='array'),
            ],
        )

        self.assertEqual(parsed.document_mode, 'form_layout_mode')
        self.assertIsNotNone(parsed.form_model)
        self.assertTrue(any(group.group_id == 'tax_residency' for group in parsed.form_model.groups))
        self.assertEqual(
            columns,
            [
                'organizationName',
                'innOrKio',
                'isResidentRF',
                'isTaxResidencyOnlyRF',
                'fatcaBeneficiaryOptionList',
            ],
        )
        self.assertEqual(rows[0]['organizationName'], 'ООО "Рога и копыта"')
        self.assertEqual(rows[0]['innOrKio'], '1234567890')
        self.assertEqual(rows[0]['isResidentRF'], 'NOWHERE')
        self.assertEqual(rows[0]['isTaxResidencyOnlyRF'], 'NO')
        self.assertEqual(rows[0]['fatcaBeneficiaryOptionList'], ['IS_FATCA_FOREIGN_INSTITUTE'])
        self.assertIn('Generated mapping from form-aware extraction.', warnings)
        self.assertEqual(parsed.form_model.layout_meta['final_source_mode'], 'form_resolver')
        self.assertEqual(parsed.form_model.layout_meta['pipeline_layers']['layout_extraction']['status'], 'completed')
        self.assertEqual(parsed.form_model.layout_meta['pipeline_layers']['generic_form_understanding']['status'], 'completed')
        self.assertEqual(parsed.form_model.layout_meta['pipeline_layers']['business_mapping']['status'], 'completed')
        resolution = next(item for item in parsed.form_model.resolved_fields if item.field == 'organizationName')
        self.assertEqual(resolution.resolved_by, 'form_resolver')

    def test_docx_form_like_single_choice_ambiguity_is_not_silently_resolved(self) -> None:
        path = self.test_root / 'form_layout_ambiguous.docx'
        doc = Document()
        table = doc.add_table(rows=4, cols=2)
        table.rows[0].cells[0].text = 'Является ли выгодоприобретатель налоговым резидентом только в РФ'
        table.rows[1].cells[0].text = 'X'
        table.rows[1].cells[1].text = 'ДА'
        table.rows[2].cells[0].text = 'X'
        table.rows[2].cells[1].text = 'Не являюсь налоговым резидентом ни в одном государстве'
        table.rows[3].cells[0].text = 'ИНН/КИО'
        table.rows[3].cells[1].text = '1234567890'
        doc.save(path)

        parsed = parse_file(path, path.name)
        columns, rows, warnings = resolve_generation_source(
            parsed,
            target_fields=[
                TargetField(name='isResidentRF', type='string'),
                TargetField(name='innOrKio', type='string'),
            ],
        )

        self.assertEqual(parsed.document_mode, 'form_layout_mode')
        self.assertEqual(columns, ['innOrKio'])
        self.assertEqual(rows, [{'innOrKio': '1234567890'}])
        self.assertTrue(any('ambiguous' in warning.lower() for warning in warnings))
        self.assertIsNotNone(parsed.form_model)
        resolution = next(item for item in parsed.form_model.resolved_fields if item.field == 'isResidentRF')
        self.assertEqual(parsed.form_model.layout_meta['final_source_mode'], 'fallback_blocked')
        self.assertEqual(resolution.status, 'ambiguous')
        self.assertEqual(resolution.resolved_by, 'fallback_blocked')
        self.assertIsNone(resolution.value)

    def test_form_layout_without_target_fields_prefers_generic_form_source(self) -> None:
        path = self.test_root / 'form_layout_generic_source.docx'
        doc = Document()
        table = doc.add_table(rows=8, cols=2)
        table.rows[0].cells[0].text = 'РќР°РёРјРµРЅРѕРІР°РЅРёРµ РѕСЂРіР°РЅРёР·Р°С†РёРё'
        table.rows[0].cells[1].text = 'РћРћРћ "Р РѕРіР° Рё РєРѕРїС‹С‚Р°"'
        table.rows[1].cells[0].text = 'РРќРќ/РљРРћ'
        table.rows[1].cells[1].text = '1234567890'
        table.rows[2].cells[0].text = 'РЇРІР»СЏРµС‚СЃСЏ Р»Рё РІС‹РіРѕРґРѕРїСЂРёРѕР±СЂРµС‚Р°С‚РµР»СЊ РЅР°Р»РѕРіРѕРІС‹Рј СЂРµР·РёРґРµРЅС‚РѕРј С‚РѕР»СЊРєРѕ РІ Р Р¤'
        table.rows[3].cells[0].text = ' '
        table.rows[3].cells[1].text = 'Р”Рђ'
        table.rows[4].cells[0].text = 'X'
        table.rows[4].cells[1].text = 'РќРµ СЏРІР»СЏСЋСЃСЊ РЅР°Р»РѕРіРѕРІС‹Рј СЂРµР·РёРґРµРЅС‚РѕРј РЅРё РІ РѕРґРЅРѕРј РіРѕСЃСѓРґР°СЂСЃС‚РІРµ'
        table.rows[5].cells[0].text = 'FATCA СЃС‚Р°С‚СѓСЃ РІС‹РіРѕРґРѕРїСЂРёРѕР±СЂРµС‚Р°С‚РµР»СЏ'
        table.rows[6].cells[0].text = 'X'
        table.rows[6].cells[1].text = 'РРЅРѕСЃС‚СЂР°РЅРЅС‹Рј С„РёРЅР°РЅСЃРѕРІС‹Рј РёРЅСЃС‚РёС‚СѓС‚РѕРј'
        table.rows[7].cells[0].text = ' '
        table.rows[7].cells[1].text = 'Р‘РѕР»РµРµ 10% Р°РєС†РёР№ РїСЂРёРЅР°РґР»РµР¶Р°С‚ РЅР°Р»РѕРіРѕРїР»Р°С‚РµР»СЊС‰РёРєР°Рј РЎРЁРђ'
        doc.save(path)

        parsed = parse_file(path, path.name)
        columns, rows, warnings = resolve_generation_source(parsed)

        self.assertEqual(parsed.content_type, 'form')
        self.assertEqual(parsed.document_mode, 'form_layout_mode')
        self.assertEqual(parsed.form_model.layout_meta['final_source_mode'], 'generic_form_source')
        self.assertGreaterEqual(len(columns), 2)
        self.assertEqual(len(rows), 1)
        self.assertIn('fatca_beneficiary', columns)
        self.assertTrue(any(value == '1234567890' for value in rows[0].values()))
        self.assertIsInstance(rows[0]['fatca_beneficiary'], list)
        self.assertIn('Generated mapping from form-aware extracted fields.', warnings)

    def test_docx_question_option_row_is_not_treated_as_scalar(self) -> None:
        scalar = _extract_scalar_from_table_row(
            ['Российской Федерации?', 'X Не являюсь налоговым резидентом ни в одном государстве (территории)'],
            table_index=0,
            row_index=4,
        )

        self.assertIsNone(scalar)

    def test_layout_lines_extract_docx_table_scalars_without_raw_form_rows(self) -> None:
        scalars = _extract_scalars_from_layout_lines(
            [
                {
                    'text': 'Наименование организации',
                    'table_idx': 0,
                    'row_idx': 1,
                    'cell_idx': 0,
                    'source_type': 'table_cell',
                },
                {
                    'text': 'ООО «Рога и Копыта»',
                    'table_idx': 0,
                    'row_idx': 1,
                    'cell_idx': 1,
                    'source_type': 'table_cell',
                },
                {
                    'text': 'ИНН/КИО',
                    'table_idx': 0,
                    'row_idx': 2,
                    'cell_idx': 0,
                    'source_type': 'table_cell',
                },
                {
                    'text': '1234567890',
                    'table_idx': 0,
                    'row_idx': 2,
                    'cell_idx': 1,
                    'source_type': 'table_cell',
                },
                {
                    'text': 'Российской Федерации?',
                    'table_idx': 0,
                    'row_idx': 4,
                    'cell_idx': 0,
                    'source_type': 'table_cell',
                },
                {
                    'text': 'X Не являюсь налоговым резидентом ни в одном государстве (территории)',
                    'table_idx': 0,
                    'row_idx': 4,
                    'cell_idx': 1,
                    'source_type': 'table_cell',
                },
            ]
        )

        pairs = {(scalar['label'], scalar['value']) for scalar in scalars}
        self.assertIn(('Наименование организации', 'ООО «Рога и Копыта»'), pairs)
        self.assertIn(('ИНН/КИО', '1234567890'), pairs)
        self.assertNotIn(
            ('Российской Федерации?', 'X Не являюсь налоговым резидентом ни в одном государстве (территории)'),
            pairs,
        )

    def test_parse_kv_line_rejects_verbose_consent_clause(self) -> None:
        scalar = _parse_kv_line(
            'настоящее согласие предоставляется на совершение следующих действий с персональными данными: '
            'передача (в том числе трансграничная), сбор, запись, систематизация, накопление, хранение, '
            'уточнение (обновление, изменение), извлечение, использование, обезличивание, блокирование, '
            'удаление, уничтожение.',
            source_ref={'source_type': 'paragraph', 'paragraph_idx': 1},
        )

        self.assertIsNone(scalar)

    def test_extract_scalar_from_table_row_rejects_heading_fragment_pair(self) -> None:
        scalar = _extract_scalar_from_table_row(
            ['СВЕДЕНИЯ О ВЫГОДОПРИОБРЕТАТЕЛЕ', 'ЮРИДИЧЕСКОМ ЛИЦЕ или ИНОСТРАННОЙ СТРУКТУРЕ БЕЗ'],
            table_index=0,
            row_index=1,
        )

        self.assertIsNone(scalar)

    def test_understand_generic_form_suppresses_scalars_consumed_by_groups(self) -> None:
        form_model = understand_generic_form(
            layout_layer={
                'seed_scalars': [
                    {
                        'label': 'Наименование организации',
                        'value': 'ООО «Рога и Копыта»',
                        'source_ref': {'table_idx': 0, 'row_idx': 1, 'source_type': 'table_cell'},
                        'confidence': 'high',
                    },
                    {
                        'label': 'ДА, является налоговым резидентом только в РФ',
                        'value': 'ДА, является налоговым резидентом только в РФ',
                        'source_ref': {'table_idx': 0, 'row_idx': 3, 'source_type': 'table_cell'},
                        'confidence': 'medium',
                    },
                    {
                        'label': 'НЕТ, является налоговым резидентом в следующем(их) иностранном(ых)',
                        'value': 'НЕТ, является налоговым резидентом в следующем(их) иностранном(ых)',
                        'source_ref': {'table_idx': 0, 'row_idx': 5, 'source_type': 'table_cell'},
                        'confidence': 'medium',
                    },
                ],
                'layout_lines': [],
                'raw_table_rows': [
                    {
                        'table_idx': 0,
                        'row_idx': 2,
                        'cells': ['Является ли выгодоприобретатель налоговым резидентом только в Российской Федерации?', ''],
                        'cell_paragraphs': [['Является ли выгодоприобретатель налоговым резидентом только в Российской Федерации?'], []],
                    },
                    {
                        'table_idx': 0,
                        'row_idx': 3,
                        'cells': ['', 'ДА, является налоговым резидентом только в РФ'],
                        'cell_paragraphs': [[], ['ДА, является налоговым резидентом только в РФ']],
                    },
                    {
                        'table_idx': 0,
                        'row_idx': 4,
                        'cells': ['X', 'Не являюсь налоговым резидентом ни в одном государстве (территории)'],
                        'cell_paragraphs': [['X'], ['Не являюсь налоговым резидентом ни в одном государстве (территории)']],
                    },
                    {
                        'table_idx': 0,
                        'row_idx': 5,
                        'cells': ['', 'НЕТ, является налоговым резидентом в следующем(их) иностранном(ых)'],
                        'cell_paragraphs': [[], ['НЕТ, является налоговым резидентом в следующем(их) иностранном(ых)']],
                    },
                ],
                'layout_meta': {'pipeline_layers': {'layout_extraction': {'status': 'completed'}}},
                'sections': [],
            },
            tables=[],
            kv_pairs=[],
        )

        self.assertIsNotNone(form_model)
        scalar_pairs = {(scalar['label'], scalar['value']) for scalar in form_model['scalars']}
        self.assertIn(('Наименование организации', 'ООО «Рога и Копыта»'), scalar_pairs)
        self.assertNotIn(
            ('ДА, является налоговым резидентом только в РФ', 'ДА, является налоговым резидентом только в РФ'),
            scalar_pairs,
        )
        self.assertNotIn(
            ('НЕТ, является налоговым резидентом в следующем(их) иностранном(ых)', 'НЕТ, является налоговым резидентом в следующем(их) иностранном(ых)'),
            scalar_pairs,
        )
        self.assertTrue(any(group['group_id'] == 'tax_residency' for group in form_model['groups']))

    def test_understand_generic_form_filters_heading_like_seed_scalar(self) -> None:
        form_model = understand_generic_form(
            layout_layer={
                'seed_scalars': [
                    {
                        'label': 'СВЕДЕНИЯ О ВЫГОДОПРИОБРЕТАТЕЛЕ',
                        'value': 'ЮРИДИЧЕСКОМ ЛИЦЕ или ИНОСТРАННОЙ СТРУКТУРЕ БЕЗ',
                        'source_ref': {'source_type': 'line', 'line_id': 'line-1'},
                        'confidence': 'medium',
                    },
                    {
                        'label': 'Наименование организации',
                        'value': 'ООО «Рога и Копыта»',
                        'source_ref': {'source_type': 'line', 'line_id': 'line-2'},
                        'confidence': 'medium',
                    },
                ],
                'layout_lines': [],
                'raw_table_rows': [],
                'layout_meta': {'pipeline_layers': {'layout_extraction': {'status': 'completed'}}},
                'sections': [],
            },
            tables=[],
            kv_pairs=[],
        )

        self.assertIsNotNone(form_model)
        scalar_pairs = {(scalar['label'], scalar['value']) for scalar in form_model['scalars']}
        self.assertIn(('Наименование организации', 'ООО «Рога и Копыта»'), scalar_pairs)
        self.assertNotIn(
            ('СВЕДЕНИЯ О ВЫГОДОПРИОБРЕТАТЕЛЕ', 'ЮРИДИЧЕСКОМ ЛИЦЕ или ИНОСТРАННОЙ СТРУКТУРЕ БЕЗ'),
            scalar_pairs,
        )

    def test_generic_form_source_prefers_form_model_over_dirty_raw_rows(self) -> None:
        parsed = ParsedFile(
            file_name='dirty_form_rows.docx',
            file_type='docx',
            columns=['column1', 'column2'],
            rows=[
                {'column1': 'ООО «Рога и Копыта»', 'column2': 'ООО «Рога и Копыта»'},
                {'column1': '1234567890', 'column2': '1234567890'},
                {'column1': 'ДА, является налоговым резидентом только в РФ', 'column2': 'ДА, является налоговым резидентом только в РФ'},
                {'column1': 'НЕТ, является налоговым резидентом в следующем(их) иностранном(ых)', 'column2': 'НЕТ, является налоговым резидентом в следующем(их) иностранном(ых)'},
            ],
            content_type='form',
            document_mode='form_layout_mode',
            extraction_status='text_extracted',
            raw_text='',
            text_blocks=[],
            sections=[],
            kv_pairs=[],
            source_candidates=[],
            sheets=[],
            form_model=_build_form_model(
                {
                    'scalars': [
                        {
                            'label': 'Наименование организации',
                            'value': 'ООО «Рога и Копыта»',
                            'source_ref': {'table_idx': 0, 'row_idx': 1},
                            'confidence': 'high',
                        },
                        {
                            'label': 'ИНН/КИО',
                            'value': '1234567890',
                            'source_ref': {'table_idx': 0, 'row_idx': 2},
                            'confidence': 'high',
                        },
                    ],
                    'groups': [
                        {
                            'group_id': 'tax_residency',
                            'question': 'Является ли выгодоприобретатель налоговым резидентом только в Российской Федерации?',
                            'group_type': 'single_choice',
                            'options': [
                                {
                                    'label': 'ДА, является налоговым резидентом только в РФ',
                                    'selected': False,
                                    'marker_text': '',
                                    'source_ref': {'table_idx': 0, 'row_idx': 3},
                                },
                                {
                                    'label': 'Не являюсь налоговым резидентом ни в одном государстве (территории)',
                                    'selected': True,
                                    'marker_text': 'X',
                                    'source_ref': {'table_idx': 0, 'row_idx': 4},
                                },
                                {
                                    'label': 'НЕТ, является налоговым резидентом в следующем(их) иностранном(ых)',
                                    'selected': False,
                                    'marker_text': '',
                                    'source_ref': {'table_idx': 0, 'row_idx': 5},
                                },
                            ],
                            'source_ref': {'table_idx': 0, 'row_idx': 3},
                        }
                    ],
                    'layout_lines': [],
                    'layout_meta': {},
                    'resolved_fields': [],
                }
            ),
            warnings=[],
        )

        columns, rows, warnings = _resolve_generic_form_layout_source(parsed)

        self.assertEqual(columns, ['Наименование организации', 'ИНН/КИО', 'tax_residency'])
        self.assertEqual(
            rows,
            [
                {
                    'Наименование организации': 'ООО «Рога и Копыта»',
                    'ИНН/КИО': '1234567890',
                    'tax_residency': 'Не являюсь налоговым резидентом ни в одном государстве (территории)',
                }
            ],
        )
        self.assertNotIn('ООО «Рога и Копыта»', columns)
        self.assertNotIn('1234567890', columns)
        self.assertIn('Generated mapping from form-aware extracted fields.', warnings)

    def test_docx_inline_cell_paragraphs_build_fatca_group(self) -> None:
        _scalars, groups = _extract_table_rows_as_form(
            [
                {
                    'table_idx': 0,
                    'row_idx': 11,
                    'cells': [
                        'Является ли хотя бы одно из следующих утверждений для выгодоприобретателя верным:',
                        'Являюсь лицом, неотделимым от собственника для целей налогообложения в США (disregarded entity);',
                    ],
                    'cell_paragraphs': [
                        ['Является ли хотя бы одно из следующих утверждений для выгодоприобретателя верным:'],
                        [
                            'Являюсь лицом, неотделимым от собственника для целей налогообложения в США (disregarded entity);',
                            'Собственник (owner) disregarded entity является',
                            'X Являюсь Иностранным финансовым институтом для целей FATCA;',
                            'Более 10% акций (долей) принадлежат налогоплательщикам США (юр. лицам/физ.лицам).',
                            'НЕТ, данные утверждения не применимы для организации',
                        ],
                    ],
                }
            ],
            table_index=0,
        )

        self.assertEqual(len(groups), 1)
        group = groups[0]
        self.assertEqual(group['group_id'], 'fatca_beneficiary')
        self.assertTrue(any(option['selected'] for option in group['options']))
        selected = [option['label'] for option in group['options'] if option['selected']]
        self.assertEqual(selected, ['Являюсь Иностранным финансовым институтом для целей FATCA'])
        resolution = _resolve_fatca_group('fatcaBeneficiaryOptionList', groups, layout_lines=[])
        self.assertEqual(resolution['status'], 'resolved')
        self.assertEqual(resolution['value'], ['IS_FATCA_FOREIGN_INSTITUTE'])

    def test_form_critical_field_does_not_fall_back_to_legacy_candidates(self) -> None:
        path = self.test_root / 'critical_fallback_blocked.txt'
        path.write_text(
            'Анкета\n\nИНН/КИО: 1234567890\nФИО: Иванов Иван\n',
            encoding='utf-8',
        )

        parsed = parse_file(path, path.name)
        columns, rows, warnings = resolve_generation_source(
            parsed,
            target_fields=[TargetField(name='isResidentRF', type='string')],
        )

        self.assertEqual(parsed.document_mode, 'form_layout_mode')
        self.assertEqual(columns, [])
        self.assertEqual(rows, [])
        self.assertEqual(parsed.form_model.layout_meta['final_source_mode'], 'fallback_blocked')
        self.assertTrue(any('blocked' in warning.lower() for warning in warnings))

    def test_form_non_critical_request_can_use_legacy_fallback_with_provenance(self) -> None:
        path = self.test_root / 'non_critical_fallback.txt'
        path.write_text(
            'Анкета\n\nФИО: Иванов Иван\nДата рождения: 01.01.1990\n',
            encoding='utf-8',
        )

        parsed = parse_file(path, path.name)
        columns, rows, warnings = resolve_generation_source(
            parsed,
            target_fields=[TargetField(name='customerNameNormalized', type='string')],
        )

        self.assertEqual(parsed.document_mode, 'form_layout_mode')
        self.assertEqual(parsed.form_model.layout_meta['final_source_mode'], 'legacy_fallback')
        self.assertEqual(columns, ['ФИО', 'Дата рождения'])
        self.assertEqual(rows[0]['ФИО'], 'Иванов Иван')
        self.assertTrue(any('fell back' in warning.lower() for warning in warnings))


class PdfAndRepairParserTests(unittest.TestCase):
    def test_pdf_region_zoning_separates_text_and_noise(self) -> None:
        zone_summary = classify_pdf_document_zones(
            tables=[],
            layout_lines=[
                {
                    'line_id': 'line-1',
                    'text': 'СВЕДЕНИЯ О ВЫГОДОПРИОБРЕТАТЕЛЕ',
                    'page': 1,
                    'column_id': 1,
                    'y': 10.0,
                },
                {
                    'line_id': 'line-2',
                    'text': 'Выгодоприобретатель ведет операционную деятельность на территории нескольких государств.',
                    'page': 1,
                    'column_id': 1,
                    'y': 42.0,
                },
                {
                    'line_id': 'line-3',
                    'text': 'Организация использует международные расчеты и обслуживает внешнеторговые контракты.',
                    'page': 1,
                    'column_id': 1,
                    'y': 56.0,
                },
            ],
            raw_text='',
        )

        self.assertGreaterEqual(len(zone_summary['region_zones']), 2)
        self.assertTrue(any(zone['zone_type'] == 'noise' for zone in zone_summary['region_zones']))
        self.assertTrue(any(zone['zone_type'] == 'text' for zone in zone_summary['region_zones']))
        self.assertIn('zone_graph', zone_summary)
        self.assertIn('parser_outputs', zone_summary)
        self.assertTrue(any(node['zone_confidence'] is not None for node in zone_summary['zone_graph']['nodes']))
        self.assertGreaterEqual(len(zone_summary['parser_outputs']['text']['regions']), 1)
        self.assertGreaterEqual(len(zone_summary['parser_outputs']['noise']['regions']), 1)

    def test_pdf_form_like_table_routes_to_form_parser(self) -> None:
        fake_pdf = FakePdfContext(
            [
                FakePdfPage(
                    text='Form-like PDF with table layout and checkbox options',
                    tables=[
                        [
                            ['Является ли выгодоприобретатель налоговым резидентом только в РФ', ''],
                            ['', 'ДА'],
                            ['X', 'Не являюсь налоговым резидентом ни в одном государстве'],
                            ['', 'НЕТ, является налоговым резидентом в иностранном государстве'],
                        ]
                    ],
                )
            ]
        )

        with patch('pdf_parser.pdfplumber.open', return_value=fake_pdf):
            parsed = parse_file(Path('pdf_form_table.pdf'), 'pdf_form_table.pdf')
            columns, rows, _warnings = resolve_generation_source(
                parsed,
                target_fields=[TargetField(name='isResidentRF', type='string')],
            )

        self.assertEqual(parsed.content_type, 'form')
        self.assertEqual(parsed.document_mode, 'form_layout_mode')
        self.assertEqual(columns, ['isResidentRF'])
        self.assertEqual(rows[0]['isResidentRF'], 'NOWHERE')
        self.assertTrue(any('PDF zoning summary' in warning for warning in parsed.warnings))

    def test_pdf_regular_table_routes_to_data_table_parser(self) -> None:
        fake_pdf = FakePdfContext(
            [
                FakePdfPage(
                    text='Simple tabular PDF',
                    tables=[
                        [
                            ['name', 'amount'],
                            ['Alice', '10'],
                            ['Bob', '20'],
                        ]
                    ],
                )
            ]
        )

        with patch('pdf_parser.pdfplumber.open', return_value=fake_pdf):
            parsed = parse_file(Path('pdf_data_table.pdf'), 'pdf_data_table.pdf')
            columns, rows, warnings = resolve_generation_source(parsed)

        self.assertEqual(parsed.content_type, 'table')
        self.assertEqual(parsed.document_mode, 'data_table_mode')
        self.assertEqual(parsed.columns, ['name', 'amount'])
        self.assertEqual(parsed.rows, [{'name': 'Alice', 'amount': '10'}, {'name': 'Bob', 'amount': '20'}])
        self.assertEqual(columns, ['name', 'amount'])
        self.assertEqual(rows, [{'name': 'Alice', 'amount': '10'}, {'name': 'Bob', 'amount': '20'}])
        self.assertTrue(any('PDF zoning classified 1 table zone' in warning for warning in parsed.warnings))
        self.assertEqual(warnings, [])
        self.assertTrue(parsed.pdf_zone_summary)
        self.assertTrue(parsed.pdf_zone_summary.get('parser_outputs'))

    def test_pdf_noise_regions_do_not_override_text_content_type(self) -> None:
        fake_pdf = FakePdfContext(
            [
                FakePdfPage(
                    text='СВЕДЕНИЯ О ВЫГОДОПРИОБРЕТАТЕЛЕ\n\nНастоящее согласие предоставляется на совершение следующих действий.\nБанк вправе осуществлять обработку персональных данных.',
                    words=[
                        {'text': 'СВЕДЕНИЯ', 'x0': 20, 'x1': 86, 'top': 10, 'bottom': 18},
                        {'text': 'О', 'x0': 90, 'x1': 98, 'top': 10, 'bottom': 18},
                        {'text': 'ВЫГОДОПРИОБРЕТАТЕЛЕ', 'x0': 102, 'x1': 246, 'top': 10, 'bottom': 18},
                        {'text': 'Настоящее', 'x0': 20, 'x1': 88, 'top': 50, 'bottom': 58},
                        {'text': 'согласие', 'x0': 92, 'x1': 150, 'top': 50, 'bottom': 58},
                        {'text': 'предоставляется', 'x0': 154, 'x1': 252, 'top': 50, 'bottom': 58},
                        {'text': 'на', 'x0': 256, 'x1': 270, 'top': 50, 'bottom': 58},
                        {'text': 'совершение', 'x0': 274, 'x1': 348, 'top': 50, 'bottom': 58},
                        {'text': 'следующих', 'x0': 352, 'x1': 426, 'top': 50, 'bottom': 58},
                        {'text': 'действий.', 'x0': 430, 'x1': 494, 'top': 50, 'bottom': 58},
                        {'text': 'Банк', 'x0': 20, 'x1': 54, 'top': 64, 'bottom': 72},
                        {'text': 'вправе', 'x0': 58, 'x1': 104, 'top': 64, 'bottom': 72},
                        {'text': 'осуществлять', 'x0': 108, 'x1': 198, 'top': 64, 'bottom': 72},
                        {'text': 'обработку', 'x0': 202, 'x1': 270, 'top': 64, 'bottom': 72},
                        {'text': 'персональных', 'x0': 274, 'x1': 368, 'top': 64, 'bottom': 72},
                        {'text': 'данных.', 'x0': 372, 'x1': 426, 'top': 64, 'bottom': 72},
                    ],
                )
            ]
        )

        with patch('pdf_parser.pdfplumber.open', return_value=fake_pdf):
            parsed = parse_file(Path('pdf_text_regions.pdf'), 'pdf_text_regions.pdf')

        self.assertEqual(parsed.content_type, 'text')
        self.assertEqual(parsed.document_mode, 'data_table_mode')
        self.assertTrue(any('PDF zoning regions' in warning for warning in parsed.warnings))

    def test_image_based_pdf_can_use_external_ocr_service(self) -> None:
        fake_pdf = FakePdfContext(
            [
                FakePdfPage(
                    text='',
                    tables=[],
                    words=[],
                )
            ]
        )

        with patch('pdf_parser.pdfplumber.open', return_value=fake_pdf), patch(
            'pdf_parser.extract_text_from_ocr_service',
            return_value={
                'text': 'Наименование организации: ООО "Рога и копыта"\nИНН/КИО: 1234567890',
                'blocks': [
                    {'id': 'ocr-line-1', 'kind': 'line', 'text': 'Наименование организации: ООО "Рога и копыта"', 'label': 'ocr'},
                    {'id': 'ocr-line-2', 'kind': 'line', 'text': 'ИНН/КИО: 1234567890', 'label': 'ocr'},
                ],
                'warnings': ['Text was extracted via the external OCR service.'],
            },
        ):
            parsed = parse_file(Path('scan_pdf.pdf'), 'scan_pdf.pdf')
            columns, rows, warnings = resolve_generation_source(parsed)

        self.assertEqual(parsed.content_type, 'form')
        self.assertEqual(parsed.extraction_status, 'text_extracted')
        self.assertEqual(columns, ['Наименование организации', 'ИНН/КИО'])
        self.assertEqual(rows[0]['Наименование организации'], 'ООО "Рога и копыта"')
        self.assertTrue(any('OCR fallback' in warning for warning in parsed.warnings))

    def test_ocr_image_zone_classification_filters_noise_and_tracks_merge_stats(self) -> None:
        merged_text, scored_blocks, zone_summary = _merge_ocr_image_results(
            raw_text=(
                'СВЕДЕНИЯ О ВЫГОДОПРИОБРЕТАТЕЛЕ\n'
                'Настоящее согласие предоставляется на обработку персональных данных\n'
                'Наименование организации: ООО "Рога и копыта"\n'
                'ИНН/КИО: 1234567890'
            ),
            layout_blocks=[
                {'id': 'line-1', 'text': 'СВЕДЕНИЯ О ВЫГОДОПРИОБРЕТАТЕЛЕ', 'page': 1, 'x': 10, 'y': 10, 'width': 240, 'height': 12, 'confidence': 0.96},
                {'id': 'line-2', 'text': 'Настоящее согласие предоставляется на обработку персональных данных', 'page': 1, 'x': 10, 'y': 32, 'width': 410, 'height': 12, 'confidence': 0.88},
                {'id': 'line-3', 'text': 'Наименование организации: ООО "Рога и копыта"', 'page': 1, 'x': 10, 'y': 92, 'width': 340, 'height': 12, 'confidence': 0.94},
                {'id': 'line-4', 'text': 'ИНН/КИО: 1234567890', 'page': 1, 'x': 10, 'y': 108, 'width': 190, 'height': 12, 'confidence': 0.93},
            ],
        )

        self.assertIn('Наименование организации', merged_text)
        self.assertIn('ИНН/КИО', merged_text)
        self.assertNotIn('Настоящее согласие', merged_text)
        self.assertTrue(any(block['ocr_zone_type'] == 'noise' for block in scored_blocks))
        self.assertGreaterEqual(int(zone_summary['counts']['form']), 1)
        self.assertGreaterEqual(int(zone_summary['counts']['noise']), 1)
        self.assertGreaterEqual(int(zone_summary['merge_stats']['dropped_noise_lines']), 1)
        self.assertTrue(zone_summary['selected_region_ids'])

    def test_image_ocr_checkbox_alias_is_detected_as_selected_option(self) -> None:
        with patch(
            'document_parser.extract_text_from_ocr_service',
            return_value={
                'text': (
                    'Является ли выгодоприобретатель налоговым резидентом только в РФ\n'
                    'I\n'
                    'Не являюсь налоговым резидентом ни в одном государстве'
                ),
                'blocks': [
                    {'id': 'ocr-line-1', 'kind': 'line', 'text': 'Является ли выгодоприобретатель налоговым резидентом только в РФ', 'label': 'ocr', 'page': 1, 'x': 10, 'y': 10, 'width': 320, 'height': 10, 'confidence': 0.92},
                    {'id': 'ocr-line-2', 'kind': 'line', 'text': 'I', 'label': 'ocr', 'page': 1, 'x': 12, 'y': 34, 'width': 8, 'height': 10, 'confidence': 0.67},
                    {'id': 'ocr-line-3', 'kind': 'line', 'text': 'Не являюсь налоговым резидентом ни в одном государстве', 'label': 'ocr', 'page': 1, 'x': 34, 'y': 34, 'width': 290, 'height': 10, 'confidence': 0.9},
                ],
                'warnings': ['Text was extracted via the external OCR service.'],
                'ocr_metadata': {'engine': 'paddleocr'},
            },
        ):
            parsed = parse_file(Path('photo_tax.jpg'), 'photo_tax.jpg')
            columns, rows, _warnings = resolve_generation_source(
                parsed,
                target_fields=[TargetField(name='isResidentRF', type='string')],
            )

        self.assertTrue(parsed.ocr_used)
        self.assertEqual(parsed.document_mode, 'form_layout_mode')
        self.assertEqual(columns, ['isResidentRF'])
        self.assertEqual(rows[0]['isResidentRF'], 'NOWHERE')
        self.assertTrue(parsed.ocr_metadata.get('zone_summary'))
        tax_group = next(group for group in parsed.form_model.groups if group.group_id == 'tax_residency')
        self.assertTrue(any(option.selected for option in tax_group.options))

    def test_image_ocr_table_reconstruction_repairs_numbers_dates_and_names(self) -> None:
        with patch(
            'document_parser.extract_text_from_ocr_service',
            return_value={
                'text': (
                    'name | amount | date\n'
                    'Aнна Иванова | 12оооо | 2025-01-02\n'
                    'Иван Иванов | 5O0 | 2025-91-02'
                ),
                'blocks': [
                    {'id': 'ocr-row-1', 'kind': 'line', 'text': 'name | amount | date', 'label': 'ocr', 'page': 1, 'x': 10, 'y': 10, 'width': 220, 'height': 10, 'confidence': 0.97},
                    {'id': 'ocr-row-2', 'kind': 'line', 'text': 'Aнна Иванова | 12оооо | 2025-01-02', 'label': 'ocr', 'page': 1, 'x': 10, 'y': 28, 'width': 280, 'height': 10, 'confidence': 0.94},
                    {'id': 'ocr-row-3', 'kind': 'line', 'text': 'Иван Иванов | 5O0 | 2025-91-02', 'label': 'ocr', 'page': 1, 'x': 10, 'y': 44, 'width': 260, 'height': 10, 'confidence': 0.92},
                    {'id': 'ocr-noise-1', 'kind': 'line', 'text': 'тг', 'label': 'ocr', 'page': 1, 'x': 10, 'y': 70, 'width': 16, 'height': 10, 'confidence': 0.31},
                ],
                'warnings': ['Text was extracted via the external OCR service.'],
                'ocr_metadata': {'engine': 'paddleocr'},
            },
        ):
            parsed = parse_file(Path('ocr_table.png'), 'ocr_table.png')

        self.assertTrue(parsed.ocr_used)
        self.assertEqual(parsed.extraction_status, 'structured_extracted')
        self.assertEqual(parsed.columns, ['name', 'amount', 'date'])
        self.assertEqual(len(parsed.rows), 2)
        self.assertEqual(parsed.rows[0]['name'], 'Анна Иванова')
        self.assertEqual(parsed.rows[0]['amount'], '120000')
        self.assertEqual(parsed.rows[0]['date'], '2025-01-02')
        self.assertEqual(parsed.rows[1]['amount'], '500')
        self.assertEqual(parsed.rows[1]['date'], '')
        self.assertEqual(parsed.ocr_metadata['table_reconstruction']['column_types']['amount'], 'number')
        self.assertEqual(parsed.ocr_metadata['table_reconstruction']['column_types']['date'], 'date')
        self.assertTrue(any('OCR table reconstruction detected' in warning for warning in parsed.warnings))

    def test_image_ocr_table_reconstruction_drops_inconsistent_row_shapes(self) -> None:
        with patch(
            'document_parser.extract_text_from_ocr_service',
            return_value={
                'text': (
                    'name | amount | date\n'
                    'Анна | 10 | 2025-01-02\n'
                    'сломанная строка | 2025-01-03\n'
                    'Борис | 20 | 2025-01-04'
                ),
                'blocks': [
                    {'id': 'ocr-row-1', 'kind': 'line', 'text': 'name | amount | date', 'label': 'ocr', 'page': 1, 'x': 10, 'y': 10, 'width': 220, 'height': 10, 'confidence': 0.97},
                    {'id': 'ocr-row-2', 'kind': 'line', 'text': 'Анна | 10 | 2025-01-02', 'label': 'ocr', 'page': 1, 'x': 10, 'y': 28, 'width': 220, 'height': 10, 'confidence': 0.94},
                    {'id': 'ocr-row-3', 'kind': 'line', 'text': 'сломанная строка | 2025-01-03', 'label': 'ocr', 'page': 1, 'x': 10, 'y': 44, 'width': 220, 'height': 10, 'confidence': 0.92},
                    {'id': 'ocr-row-4', 'kind': 'line', 'text': 'Борис | 20 | 2025-01-04', 'label': 'ocr', 'page': 1, 'x': 10, 'y': 60, 'width': 220, 'height': 10, 'confidence': 0.95},
                ],
                'warnings': ['Text was extracted via the external OCR service.'],
                'ocr_metadata': {'engine': 'paddleocr'},
            },
        ):
            parsed = parse_file(Path('ocr_table_shape.png'), 'ocr_table_shape.png')

        self.assertEqual(parsed.columns, ['name', 'amount', 'date'])
        self.assertEqual(len(parsed.rows), 2)
        self.assertEqual(parsed.rows[0]['name'], 'Анна')
        self.assertEqual(parsed.rows[1]['name'], 'Борис')

    def test_image_ocr_table_reconstruction_can_cluster_row_bands_from_fragments(self) -> None:
        with patch(
            'document_parser.extract_text_from_ocr_service',
            return_value={
                'text': 'name amount date Анна 10 2025-01-02 Борис 20 2025-01-04',
                'blocks': [
                    {'id': 'hdr-1', 'kind': 'line', 'text': 'name', 'label': 'ocr', 'page': 1, 'x': 10, 'y': 10, 'width': 40, 'height': 10, 'confidence': 0.95},
                    {'id': 'hdr-2', 'kind': 'line', 'text': 'amount', 'label': 'ocr', 'page': 1, 'x': 110, 'y': 11, 'width': 55, 'height': 10, 'confidence': 0.95},
                    {'id': 'hdr-3', 'kind': 'line', 'text': 'date', 'label': 'ocr', 'page': 1, 'x': 210, 'y': 10, 'width': 40, 'height': 10, 'confidence': 0.95},
                    {'id': 'row-1a', 'kind': 'line', 'text': 'Анна', 'label': 'ocr', 'page': 1, 'x': 10, 'y': 30, 'width': 40, 'height': 10, 'confidence': 0.93},
                    {'id': 'row-1b', 'kind': 'line', 'text': '10', 'label': 'ocr', 'page': 1, 'x': 118, 'y': 31, 'width': 18, 'height': 10, 'confidence': 0.91},
                    {'id': 'row-1c', 'kind': 'line', 'text': '2025-01-02', 'label': 'ocr', 'page': 1, 'x': 208, 'y': 30, 'width': 78, 'height': 10, 'confidence': 0.94},
                    {'id': 'row-2a', 'kind': 'line', 'text': 'Борис', 'label': 'ocr', 'page': 1, 'x': 10, 'y': 48, 'width': 48, 'height': 10, 'confidence': 0.93},
                    {'id': 'row-2b', 'kind': 'line', 'text': '20', 'label': 'ocr', 'page': 1, 'x': 118, 'y': 49, 'width': 18, 'height': 10, 'confidence': 0.91},
                    {'id': 'row-2c', 'kind': 'line', 'text': '2025-01-04', 'label': 'ocr', 'page': 1, 'x': 208, 'y': 48, 'width': 78, 'height': 10, 'confidence': 0.94},
                ],
                'warnings': ['Text was extracted via the external OCR service.'],
                'ocr_metadata': {'engine': 'paddleocr'},
            },
        ):
            parsed = parse_file(Path('ocr_table_bands.png'), 'ocr_table_bands.png')

        self.assertEqual(parsed.columns, ['name', 'amount', 'date'])
        self.assertEqual(len(parsed.rows), 2)
        self.assertEqual(parsed.rows[0]['name'], 'Анна')
        self.assertEqual(parsed.rows[0]['amount'], '10')
        self.assertEqual(parsed.rows[1]['date'], '2025-01-04')
        self.assertEqual(parsed.ocr_metadata['table_reconstruction']['recovered_table_confidence'], 'high')
        self.assertGreaterEqual(parsed.ocr_metadata['table_reconstruction']['row_stats']['accepted_rows'], 2)

    def test_llamaparse_markdown_table_is_routed_to_table_not_unknown_group(self) -> None:
        with patch(
            'document_parser.extract_text_from_llamaparse',
            return_value={
                'text': '| ПБ | по | П |\n| -- | -- | -- |\n| 6 | 4 | 13 |\n| 4 | 6 | 15 |',
                'blocks': [
                    {'id': 'lp-1', 'kind': 'line', 'text': '| ПБ | по | П |', 'label': 'llamaparse', 'page': 1},
                    {'id': 'lp-2', 'kind': 'line', 'text': '| -- | -- | -- |', 'label': 'llamaparse', 'page': 1},
                    {'id': 'lp-3', 'kind': 'line', 'text': '| 6 | 4 | 13 |', 'label': 'llamaparse', 'page': 1},
                    {'id': 'lp-4', 'kind': 'line', 'text': '| 4 | 6 | 15 |', 'label': 'llamaparse', 'page': 1},
                ],
                'warnings': ['Primary extraction used LlamaParse.'],
            },
        ):
            parsed = parse_file(Path('llamaparse_markdown_table.png'), 'llamaparse_markdown_table.png')

        self.assertEqual(parsed.content_type, 'table')
        self.assertEqual(parsed.document_mode, 'data_table_mode')
        self.assertEqual(parsed.columns, ['ПБ', 'по', 'П'])
        self.assertEqual(len(parsed.rows), 2)
        self.assertEqual(parsed.rows[0]['П'], '13')
        self.assertFalse(parsed.form_model and parsed.form_model.groups)

    def test_pdf_zone_routing_can_prefer_table_source_over_form_source(self) -> None:
        parsed = ParsedFile(
            file_name='routed.pdf',
            file_type='pdf',
            columns=['name', 'amount'],
            rows=[{'name': 'Alice', 'amount': '10'}],
            content_type='form',
            document_mode='form_layout_mode',
            extraction_status='structured_extracted',
            raw_text='',
            text_blocks=[],
            sections=[],
            kv_pairs=[],
            source_candidates=[],
            sheets=[],
            form_model=_build_form_model(
                {
                    'scalars': [{'label': 'Наименование организации', 'value': 'ООО', 'source_ref': {'line_id': 'line-1'}}],
                    'groups': [
                        {
                            'group_id': 'group_1',
                            'question': 'Question',
                            'group_type': 'single_choice',
                            'options': [{'label': 'Option A', 'selected': True, 'source_ref': {'line_id': 'line-2'}}],
                            'source_ref': {'line_id': 'line-1'},
                        }
                    ],
                    'layout_lines': [
                        {'text': 'Question', 'line_id': 'line-1'},
                        {'text': 'Option A', 'line_id': 'line-2'},
                    ],
                    'layout_meta': {},
                    'resolved_fields': [],
                }
            ),
            pdf_zone_summary={
                'dominant_zone': 'table',
                'parser_outputs': {
                    'table': {'zones': [{'zone_confidence': 0.92}]},
                    'form': {'zones': [{'zone_confidence': 0.42}]},
                    'text': {'zones': []},
                    'noise': {'zones': []},
                },
            },
            warnings=[],
        )

        columns, rows, warnings = resolve_generation_source(parsed)

        self.assertEqual(columns, ['name', 'amount'])
        self.assertEqual(rows, [{'name': 'Alice', 'amount': '10'}])
        self.assertTrue(any('preferred tabular extraction' in warning.lower() for warning in warnings))
        self.assertEqual(parsed.form_model.layout_meta['pdf_zone_routing']['prefer_table_source'], True)

    def test_pdf_layout_words_are_grouped_and_ground_tax_option(self) -> None:
        fake_pdf = FakePdfContext(
            [
                FakePdfPage(
                    text='Tax residency form with enough extracted text for parsing flow',
                    words=[
                        {'text': 'Является', 'x0': 10, 'x1': 70, 'top': 10, 'bottom': 18},
                        {'text': 'ли', 'x0': 74, 'x1': 84, 'top': 10, 'bottom': 18},
                        {'text': 'выгодоприобретатель', 'x0': 88, 'x1': 180, 'top': 10, 'bottom': 18},
                        {'text': 'налоговым', 'x0': 184, 'x1': 250, 'top': 10, 'bottom': 18},
                        {'text': 'резидентом', 'x0': 254, 'x1': 320, 'top': 10, 'bottom': 18},
                        {'text': 'только', 'x0': 324, 'x1': 365, 'top': 10, 'bottom': 18},
                        {'text': 'в', 'x0': 369, 'x1': 376, 'top': 10, 'bottom': 18},
                        {'text': 'РФ', 'x0': 380, 'x1': 395, 'top': 10, 'bottom': 18},
                        {'text': 'X', 'x0': 14, 'x1': 20, 'top': 30, 'bottom': 38},
                        {'text': 'Не', 'x0': 40, 'x1': 52, 'top': 30, 'bottom': 38},
                        {'text': 'являюсь', 'x0': 56, 'x1': 108, 'top': 30, 'bottom': 38},
                        {'text': 'налоговым', 'x0': 112, 'x1': 178, 'top': 30, 'bottom': 38},
                        {'text': 'резидентом', 'x0': 182, 'x1': 248, 'top': 30, 'bottom': 38},
                        {'text': 'ни', 'x0': 252, 'x1': 264, 'top': 30, 'bottom': 38},
                        {'text': 'в', 'x0': 268, 'x1': 275, 'top': 30, 'bottom': 38},
                        {'text': 'одном', 'x0': 279, 'x1': 316, 'top': 30, 'bottom': 38},
                        {'text': 'государстве', 'x0': 320, 'x1': 388, 'top': 30, 'bottom': 38},
                    ],
                )
            ]
        )

        with patch('pdf_parser.pdfplumber.open', return_value=fake_pdf):
            parsed = parse_file(Path('grounded.pdf'), 'grounded.pdf')
            columns, rows, _warnings = resolve_generation_source(
                parsed,
                target_fields=[TargetField(name='isResidentRF', type='string')],
            )

        self.assertEqual(parsed.document_mode, 'form_layout_mode')
        self.assertIsNotNone(parsed.form_model)
        self.assertEqual(columns, ['isResidentRF'])
        self.assertEqual(rows[0]['isResidentRF'], 'NOWHERE')
        self.assertEqual(parsed.form_model.layout_meta['final_source_mode'], 'form_resolver')
        tax_group = next(group for group in parsed.form_model.groups if group.group_id == 'tax_residency')
        self.assertTrue(any(option.selected for option in tax_group.options))

    def test_repair_mode_marks_resolution_as_repair_model(self) -> None:
        from form_layout import build_form_document_model

        text = (
            'Является ли выгодоприобретатель налоговым резидентом только в РФ\n'
            'X ДА\n'
            'X Не являюсь налоговым резидентом ни в одном государстве\n'
        )
        parsed = ParsedFile(
            file_name='repair.txt',
            file_type='txt',
            columns=[],
            rows=[],
            content_type='form',
            document_mode='form_layout_mode',
            extraction_status='text_extracted',
            raw_text=text,
            text_blocks=[],
            sections=[],
            kv_pairs=[],
            source_candidates=[],
            sheets=[],
            form_model=_build_form_model(
                build_form_document_model(
                    file_path=Path('repair.txt'),
                    file_type='txt',
                    raw_text=text,
                    kv_pairs=[],
                    text_blocks=[],
                    sections=[],
                    layout_blocks=[],
                )
            ),
            warnings=[],
        )

        with patch(
            'form_layout.suggest_form_field_repair',
            return_value=({'status': 'resolved', 'enum_value': 'NOWHERE', 'confidence': 0.77}, []),
        ):
            columns, rows, _warnings = resolve_generation_source(
                parsed,
                target_fields=[TargetField(name='isResidentRF', type='string')],
            )

        self.assertEqual(columns, ['isResidentRF'])
        self.assertEqual(rows[0]['isResidentRF'], 'NOWHERE')
        self.assertEqual(parsed.form_model.layout_meta['final_source_mode'], 'repair_model')
        resolution = next(item for item in parsed.form_model.resolved_fields if item.field == 'isResidentRF')
        self.assertEqual(resolution.resolved_by, 'repair_model')
        self.assertEqual(resolution.status, 'resolved')

    def test_simple_ambiguous_group_does_not_trigger_repair_model(self) -> None:
        from form_layout import build_form_document_model

        text = (
            'Является ли выгодоприобретатель налоговым резидентом только в РФ\n'
            'X ДА\n'
            'X НЕТ\n'
        )
        parsed = ParsedFile(
            file_name='simple_repair.txt',
            file_type='txt',
            columns=[],
            rows=[],
            content_type='form',
            document_mode='form_layout_mode',
            extraction_status='text_extracted',
            raw_text=text,
            text_blocks=[],
            sections=[],
            kv_pairs=[],
            source_candidates=[],
            sheets=[],
            form_model=_build_form_model(
                build_form_document_model(
                    file_path=Path('simple_repair.txt'),
                    file_type='txt',
                    raw_text=text,
                    kv_pairs=[],
                    text_blocks=[],
                    sections=[],
                    layout_blocks=[],
                )
            ),
            warnings=[],
        )

        with patch('form_layout.suggest_form_field_repair') as repair_mock:
            columns, rows, warnings = resolve_generation_source(
                parsed,
                target_fields=[TargetField(name='isResidentRF', type='string')],
            )

        self.assertEqual(columns, [])
        self.assertEqual(rows, [])
        repair_mock.assert_not_called()
        resolution = next(item for item in parsed.form_model.resolved_fields if item.field == 'isResidentRF')
        self.assertEqual(resolution.resolved_by, 'fallback_blocked')
        self.assertEqual(parsed.form_model.layout_meta['final_source_mode'], 'fallback_blocked')

    def test_pdf_grouping_respects_columns_and_merges_multiline_option(self) -> None:
        fake_pdf = FakePdfContext(
            [
                FakePdfPage(
                    text='Form text with columns and wrapped options for grouping',
                    words=[
                        {'text': 'Является', 'x0': 20, 'x1': 70, 'top': 10, 'bottom': 18},
                        {'text': 'ли', 'x0': 74, 'x1': 84, 'top': 10, 'bottom': 18},
                        {'text': 'выгодоприобретатель', 'x0': 88, 'x1': 180, 'top': 10, 'bottom': 18},
                        {'text': 'налоговым', 'x0': 184, 'x1': 250, 'top': 10, 'bottom': 18},
                        {'text': 'резидентом', 'x0': 254, 'x1': 320, 'top': 10, 'bottom': 18},
                        {'text': 'только', 'x0': 324, 'x1': 365, 'top': 10, 'bottom': 18},
                        {'text': 'в', 'x0': 369, 'x1': 376, 'top': 10, 'bottom': 18},
                        {'text': 'РФ', 'x0': 380, 'x1': 395, 'top': 10, 'bottom': 18},
                        {'text': 'X', 'x0': 20, 'x1': 26, 'top': 28, 'bottom': 36},
                        {'text': 'Не', 'x0': 44, 'x1': 56, 'top': 28, 'bottom': 36},
                        {'text': 'являюсь', 'x0': 60, 'x1': 112, 'top': 28, 'bottom': 36},
                        {'text': 'налоговым', 'x0': 116, 'x1': 182, 'top': 28, 'bottom': 36},
                        {'text': 'резидентом', 'x0': 186, 'x1': 252, 'top': 28, 'bottom': 36},
                        {'text': 'в', 'x0': 56, 'x1': 63, 'top': 40, 'bottom': 48},
                        {'text': 'иностранном', 'x0': 67, 'x1': 135, 'top': 40, 'bottom': 48},
                        {'text': 'государстве', 'x0': 139, 'x1': 207, 'top': 40, 'bottom': 48},
                        {'text': 'Контактное', 'x0': 470, 'x1': 535, 'top': 28, 'bottom': 36},
                        {'text': 'лицо', 'x0': 539, 'x1': 568, 'top': 28, 'bottom': 36},
                    ],
                )
            ]
        )

        with patch('pdf_parser.pdfplumber.open', return_value=fake_pdf):
            parsed = parse_file(Path('columns.pdf'), 'columns.pdf')

        self.assertIsNotNone(parsed.form_model)
        tax_group = next(group for group in parsed.form_model.groups if group.group_id == 'tax_residency')
        self.assertEqual(len(tax_group.options), 1)
        self.assertEqual(
            tax_group.options[0].label,
            'Не являюсь налоговым резидентом в иностранном государстве',
        )
        self.assertEqual(tax_group.source_ref.get('column_id'), 1)
        self.assertTrue(all(option.source_ref.get('column_id') == 1 for option in tax_group.options))

    def test_pdf_two_column_form_rows_resolve_scalars_like_docx(self) -> None:
        fake_pdf = FakePdfContext(
            [
                FakePdfPage(
                    text='Beneficiary form with enough extracted text for parsing flow and scalar reconstruction',
                    words=[
                        {'text': 'Наименование', 'x0': 20, 'x1': 100, 'top': 10, 'bottom': 18},
                        {'text': 'организации', 'x0': 104, 'x1': 184, 'top': 10, 'bottom': 18},
                        {'text': 'ООО', 'x0': 330, 'x1': 356, 'top': 10, 'bottom': 18},
                        {'text': '«Рога', 'x0': 360, 'x1': 404, 'top': 10, 'bottom': 18},
                        {'text': 'и', 'x0': 408, 'x1': 416, 'top': 10, 'bottom': 18},
                        {'text': 'Копыта»', 'x0': 420, 'x1': 476, 'top': 10, 'bottom': 18},
                        {'text': 'ИНН/КИО', 'x0': 20, 'x1': 84, 'top': 24, 'bottom': 32},
                        {'text': '1234567890', 'x0': 330, 'x1': 402, 'top': 24, 'bottom': 32},
                        {'text': 'СВЕДЕНИЯ', 'x0': 20, 'x1': 86, 'top': 38, 'bottom': 46},
                        {'text': 'О', 'x0': 90, 'x1': 98, 'top': 38, 'bottom': 46},
                        {'text': 'ВЫГОДОПРИОБРЕТАТЕЛЕ', 'x0': 102, 'x1': 246, 'top': 38, 'bottom': 46},
                        {'text': 'ЮРИДИЧЕСКОМ', 'x0': 330, 'x1': 426, 'top': 38, 'bottom': 46},
                        {'text': 'ЛИЦЕ', 'x0': 430, 'x1': 466, 'top': 38, 'bottom': 46},
                        {'text': 'или', 'x0': 470, 'x1': 490, 'top': 38, 'bottom': 46},
                        {'text': 'ИНОСТРАННОЙ', 'x0': 494, 'x1': 586, 'top': 38, 'bottom': 46},
                        {'text': 'СТРУКТУРЕ', 'x0': 590, 'x1': 670, 'top': 38, 'bottom': 46},
                        {'text': 'Является', 'x0': 20, 'x1': 70, 'top': 54, 'bottom': 62},
                        {'text': 'ли', 'x0': 74, 'x1': 84, 'top': 54, 'bottom': 62},
                        {'text': 'выгодоприобретатель', 'x0': 88, 'x1': 180, 'top': 54, 'bottom': 62},
                        {'text': 'налоговым', 'x0': 184, 'x1': 250, 'top': 54, 'bottom': 62},
                        {'text': 'резидентом', 'x0': 254, 'x1': 320, 'top': 54, 'bottom': 62},
                        {'text': 'только', 'x0': 324, 'x1': 365, 'top': 54, 'bottom': 62},
                        {'text': 'в', 'x0': 369, 'x1': 376, 'top': 54, 'bottom': 62},
                        {'text': 'РФ', 'x0': 380, 'x1': 395, 'top': 54, 'bottom': 62},
                        {'text': 'X', 'x0': 20, 'x1': 26, 'top': 70, 'bottom': 78},
                        {'text': 'Не', 'x0': 44, 'x1': 56, 'top': 70, 'bottom': 78},
                        {'text': 'являюсь', 'x0': 60, 'x1': 112, 'top': 70, 'bottom': 78},
                        {'text': 'налоговым', 'x0': 116, 'x1': 182, 'top': 70, 'bottom': 78},
                        {'text': 'резидентом', 'x0': 186, 'x1': 252, 'top': 70, 'bottom': 78},
                        {'text': 'ни', 'x0': 256, 'x1': 268, 'top': 70, 'bottom': 78},
                        {'text': 'в', 'x0': 272, 'x1': 279, 'top': 70, 'bottom': 78},
                        {'text': 'одном', 'x0': 283, 'x1': 320, 'top': 70, 'bottom': 78},
                        {'text': 'государстве', 'x0': 324, 'x1': 392, 'top': 70, 'bottom': 78},
                    ],
                )
            ]
        )

        with patch('pdf_parser.pdfplumber.open', return_value=fake_pdf):
            parsed = parse_file(Path('pdf_form_rows.pdf'), 'pdf_form_rows.pdf')
            columns, rows, _warnings = resolve_generation_source(
                parsed,
                target_fields=[
                    TargetField(name='organizationName', type='string'),
                    TargetField(name='innOrKio', type='string'),
                    TargetField(name='isResidentRF', type='string'),
                ],
            )

        self.assertEqual(parsed.document_mode, 'form_layout_mode')
        self.assertIsNotNone(parsed.form_model)
        self.assertEqual(rows[0]['organizationName'], 'ООО «Рога и Копыта»')
        self.assertEqual(rows[0]['innOrKio'], '1234567890')
        self.assertEqual(rows[0]['isResidentRF'], 'NOWHERE')
        scalar_pairs = {(scalar.label, scalar.value) for scalar in parsed.form_model.scalars}
        self.assertIn(('Наименование организации', 'ООО «Рога и Копыта»'), scalar_pairs)
        self.assertIn(('ИНН/КИО', '1234567890'), scalar_pairs)
        self.assertNotIn(
            ('СВЕДЕНИЯ О ВЫГОДОПРИОБРЕТАТЕЛЕ', 'ЮРИДИЧЕСКОМ ЛИЦЕ или ИНОСТРАННОЙ СТРУКТУРЕ'),
            scalar_pairs,
        )

    def test_pdf_cross_column_fatca_block_uses_left_question_and_right_options(self) -> None:
        fake_pdf = FakePdfContext(
            [
                FakePdfPage(
                    text='FATCA form text with enough extracted content for realistic grouping',
                    words=[
                        {'text': 'Является', 'x0': 20, 'x1': 70, 'top': 10, 'bottom': 18},
                        {'text': 'ли', 'x0': 74, 'x1': 84, 'top': 10, 'bottom': 18},
                        {'text': 'хотя', 'x0': 88, 'x1': 118, 'top': 10, 'bottom': 18},
                        {'text': 'бы', 'x0': 122, 'x1': 136, 'top': 10, 'bottom': 18},
                        {'text': 'одно', 'x0': 140, 'x1': 172, 'top': 10, 'bottom': 18},
                        {'text': 'из', 'x0': 176, 'x1': 188, 'top': 10, 'bottom': 18},
                        {'text': 'следующих', 'x0': 192, 'x1': 264, 'top': 10, 'bottom': 18},
                        {'text': 'утверждений', 'x0': 20, 'x1': 104, 'top': 24, 'bottom': 32},
                        {'text': 'для', 'x0': 108, 'x1': 128, 'top': 24, 'bottom': 32},
                        {'text': 'выгодоприобретателя', 'x0': 132, 'x1': 246, 'top': 24, 'bottom': 32},
                        {'text': 'верным:', 'x0': 20, 'x1': 78, 'top': 38, 'bottom': 46},
                        {'text': 'Являюсь', 'x0': 330, 'x1': 386, 'top': 10, 'bottom': 18},
                        {'text': 'лицом,', 'x0': 390, 'x1': 436, 'top': 10, 'bottom': 18},
                        {'text': 'неотделимым', 'x0': 440, 'x1': 528, 'top': 10, 'bottom': 18},
                        {'text': 'от', 'x0': 532, 'x1': 546, 'top': 10, 'bottom': 18},
                        {'text': 'собственника', 'x0': 550, 'x1': 642, 'top': 10, 'bottom': 18},
                        {'text': 'Собственник', 'x0': 330, 'x1': 420, 'top': 24, 'bottom': 32},
                        {'text': '(owner)', 'x0': 424, 'x1': 480, 'top': 24, 'bottom': 32},
                        {'text': 'disregarded', 'x0': 484, 'x1': 568, 'top': 24, 'bottom': 32},
                        {'text': 'entity', 'x0': 572, 'x1': 614, 'top': 24, 'bottom': 32},
                        {'text': 'является', 'x0': 618, 'x1': 684, 'top': 24, 'bottom': 32},
                        {'text': 'X', 'x0': 330, 'x1': 336, 'top': 38, 'bottom': 46},
                        {'text': 'Являюсь', 'x0': 352, 'x1': 408, 'top': 38, 'bottom': 46},
                        {'text': 'Иностранным', 'x0': 412, 'x1': 508, 'top': 38, 'bottom': 46},
                        {'text': 'финансовым', 'x0': 512, 'x1': 596, 'top': 38, 'bottom': 46},
                        {'text': 'институтом', 'x0': 600, 'x1': 680, 'top': 38, 'bottom': 46},
                        {'text': 'для', 'x0': 684, 'x1': 704, 'top': 38, 'bottom': 46},
                        {'text': 'FATCA', 'x0': 708, 'x1': 752, 'top': 38, 'bottom': 46},
                        {'text': 'Более', 'x0': 330, 'x1': 372, 'top': 52, 'bottom': 60},
                        {'text': '10%', 'x0': 376, 'x1': 404, 'top': 52, 'bottom': 60},
                        {'text': 'акций', 'x0': 408, 'x1': 448, 'top': 52, 'bottom': 60},
                        {'text': 'НЕТ,', 'x0': 330, 'x1': 362, 'top': 66, 'bottom': 74},
                        {'text': 'данные', 'x0': 366, 'x1': 418, 'top': 66, 'bottom': 74},
                        {'text': 'утверждения', 'x0': 422, 'x1': 506, 'top': 66, 'bottom': 74},
                        {'text': 'не', 'x0': 510, 'x1': 524, 'top': 66, 'bottom': 74},
                        {'text': 'применимы', 'x0': 528, 'x1': 606, 'top': 66, 'bottom': 74},
                    ],
                )
            ]
        )

        with patch('pdf_parser.pdfplumber.open', return_value=fake_pdf):
            parsed = parse_file(Path('pdf_fatca_cross_column.pdf'), 'pdf_fatca_cross_column.pdf')
            columns, rows, _warnings = resolve_generation_source(
                parsed,
                target_fields=[TargetField(name='fatcaBeneficiaryOptionList', type='array')],
            )

        self.assertEqual(parsed.document_mode, 'form_layout_mode')
        self.assertIsNotNone(parsed.form_model)
        self.assertEqual(columns, ['fatcaBeneficiaryOptionList'])
        self.assertEqual(rows[0]['fatcaBeneficiaryOptionList'], ['IS_FATCA_FOREIGN_INSTITUTE'])
        fatca_group = next(group for group in parsed.form_model.groups if 'fatca' in group.group_id or any('fatca' in option.label.lower() for option in group.options))
        self.assertIn('Является ли хотя бы одно из следующих утверждений для выгодоприобретателя верным', fatca_group.question)
        self.assertTrue(any(option.selected for option in fatca_group.options))
        self.assertFalse(str(fatca_group.question).startswith('Собственник (owner) disregarded entity является'))

    def test_form_mode_filters_consumed_option_fragments_from_candidates(self) -> None:
        form_model = {
            'groups': [
                {
                    'group_id': 'tax_residency',
                    'question': 'Является ли выгодоприобретатель налоговым резидентом только в РФ',
                    'group_type': 'single_choice',
                    'options': [
                        {'label': 'ДА, является налоговым резидентом только в РФ', 'selected': False},
                        {'label': 'Не являюсь налоговым резидентом ни в одном государстве', 'selected': True},
                        {'label': 'НЕТ, является налоговым резидентом в иностранном государстве', 'selected': False},
                    ],
                }
            ]
        }

        kv_pairs, text_facts = _suppress_consumed_group_fragments(
            kv_pairs=[
                {'label': 'Наименование организации', 'value': 'ООО "Рога и Копыта"'},
                {'label': 'X', 'value': 'Не являюсь налоговым резидентом ни в одном государстве'},
            ],
            text_facts=[
                {'label': 'ДА, является налоговым резидентом только в РФ', 'value': 'вариант'},
                {'label': 'ИНН/КИО', 'value': '1234567890'},
            ],
            form_model=form_model,
        )

        self.assertEqual(kv_pairs, [{'label': 'Наименование организации', 'value': 'ООО "Рога и Копыта"'}])
        self.assertEqual(text_facts, [{'label': 'ИНН/КИО', 'value': '1234567890'}])

    def test_generic_form_source_suppresses_option_rows_consumed_by_group(self) -> None:
        parsed = ParsedFile(
            file_name='form_pairs.txt',
            file_type='txt',
            columns=['label', 'value'],
            rows=[
                {'label': 'Наименование организации', 'value': 'ООО "Рога и Копыта"'},
                {'label': 'ДА, является налоговым резидентом только в РФ', 'value': ''},
                {'label': 'Не являюсь налоговым резидентом ни в одном государстве', 'value': 'X'},
                {'label': 'НЕТ, является налоговым резидентом в иностранном государстве', 'value': ''},
            ],
            content_type='form',
            document_mode='form_layout_mode',
            extraction_status='text_extracted',
            raw_text='',
            text_blocks=[],
            sections=[],
            kv_pairs=[],
            source_candidates=[],
            sheets=[],
            form_model=_build_form_model(
                {
                    'scalars': [],
                    'groups': [
                        {
                            'group_id': 'tax_residency',
                            'question': 'Является ли выгодоприобретатель налоговым резидентом только в РФ',
                            'group_type': 'single_choice',
                            'options': [
                                {
                                    'label': 'ДА, является налоговым резидентом только в РФ',
                                    'selected': False,
                                    'marker_text': '',
                                    'source_ref': {'table_idx': 0, 'row_idx': 2},
                                },
                                {
                                    'label': 'Не являюсь налоговым резидентом ни в одном государстве',
                                    'selected': True,
                                    'marker_text': 'X',
                                    'source_ref': {'table_idx': 0, 'row_idx': 3},
                                },
                                {
                                    'label': 'НЕТ, является налоговым резидентом в иностранном государстве',
                                    'selected': False,
                                    'marker_text': '',
                                    'source_ref': {'table_idx': 0, 'row_idx': 4},
                                },
                            ],
                            'source_ref': {'table_idx': 0, 'row_idx': 2},
                        }
                    ],
                    'layout_lines': [],
                    'layout_meta': {},
                    'resolved_fields': [],
                }
            ),
            warnings=[],
        )

        columns, rows, warnings = _resolve_generic_form_layout_source(parsed)

        self.assertEqual(columns, ['tax_residency'])
        self.assertEqual(rows[0]['tax_residency'], 'Не являюсь налоговым резидентом ни в одном государстве')
        self.assertNotIn('Наименование организации', rows[0])
        self.assertNotIn('ДА, является налоговым резидентом только в РФ', rows[0])
        self.assertNotIn('Не являюсь налоговым резидентом ни в одном государстве', rows[0])
        self.assertNotIn('НЕТ, является налоговым резидентом в иностранном государстве', rows[0])
        self.assertIn('Generated mapping from form-aware extracted fields.', warnings)

    def test_form_explainability_contains_quality_summary(self) -> None:
        parsed = ParsedFile(
            file_name='explainability.txt',
            file_type='txt',
            columns=[],
            rows=[],
            content_type='form',
            document_mode='form_layout_mode',
            extraction_status='text_extracted',
            raw_text='ИНН/КИО: 1234567890\n',
            text_blocks=[],
            sections=[],
            kv_pairs=[],
            source_candidates=[],
            sheets=[],
            form_model=_build_form_model(
                {
                    'scalars': [],
                    'groups': [
                        {
                            'group_id': 'tax_residency',
                            'question': 'Является ли выгодоприобретатель налоговым резидентом только в РФ',
                            'group_type': 'single_choice',
                            'options': [
                                {
                                    'label': 'ДА',
                                    'selected': True,
                                    'marker_text': 'X',
                                    'source_ref': {'line_id': 'line-2'},
                                },
                                {
                                    'label': 'НЕТ',
                                    'selected': True,
                                    'marker_text': 'X',
                                    'source_ref': {'line_id': 'line-3'},
                                },
                            ],
                            'source_ref': {'line_id': 'line-1'},
                        }
                    ],
                    'layout_lines': [
                        {'text': 'Является ли выгодоприобретатель налоговым резидентом только в РФ', 'line_id': 'line-1'},
                        {'text': 'X ДА', 'line_id': 'line-2'},
                        {'text': 'X НЕТ', 'line_id': 'line-3'},
                    ],
                    'layout_meta': {},
                    'resolved_fields': [
                        {
                            'field': 'isResidentRF',
                            'status': 'ambiguous',
                            'resolved_by': 'fallback_blocked',
                            'value': None,
                            'candidates': ['YES', 'NO'],
                            'source_ref': {'line_id': 'line-1'},
                            'confidence': 0.31,
                        }
                    ],
                }
            ),
            warnings=[],
        )
        parsed.form_model.layout_meta.update(
            {
                'final_source_mode': 'fallback_blocked',
                'quality_summary': {
                    'needs_attention': True,
                    'repair_recommended': True,
                    'unresolved_critical_fields': ['isResidentRF'],
                    'ambiguous_fields': ['isResidentRF'],
                    'multiple_selected_single_choice_groups': ['tax_residency'],
                    'red_flags': [{'code': 'critical_unresolved'}],
                },
                'requested_target_fields': ['isResidentRF'],
            }
        )

        explainability = _build_form_explainability(parsed)

        self.assertIsNotNone(explainability)
        self.assertEqual(explainability['final_source_mode'], 'fallback_blocked')
        self.assertTrue(explainability['quality_summary']['needs_attention'])
        self.assertEqual(explainability['quality_summary']['red_flags'][0]['code'], 'critical_unresolved')
        self.assertEqual(explainability['section_count'], 0)
        self.assertTrue(explainability['repair_plan']['recommended'])
        self.assertEqual(explainability['repair_plan']['strategy'], 'layout_chunks_then_targeted_repair')
        self.assertTrue(any(action['target_field'] == 'isResidentRF' for action in explainability['repair_plan']['actions'] if 'target_field' in action))
        repair_action = next(action for action in explainability['repair_plan']['actions'] if action.get('target_field') == 'isResidentRF')
        self.assertIn('tax_residency', repair_action['chunk_refs']['group_ids'])
        self.assertIn('line-1', repair_action['chunk_refs']['line_ids'])
        self.assertEqual(explainability['repair_plan']['llm_policy'], 'targeted_local_chunks_only')

    def test_form_explainability_builds_generic_repair_plan_without_business_quality(self) -> None:
        parsed = ParsedFile(
            file_name='generic_explainability.txt',
            file_type='txt',
            columns=[],
            rows=[],
            content_type='form',
            document_mode='form_layout_mode',
            extraction_status='text_extracted',
            raw_text='Question\nX A\nX B\n',
            text_blocks=[],
            sections=[],
            kv_pairs=[],
            source_candidates=[],
            sheets=[],
            form_model=_build_form_model(
                {
                    'scalars': [],
                    'groups': [
                        {
                            'group_id': 'group_1',
                            'question': 'Question',
                            'group_type': 'single_choice',
                            'options': [
                                {'label': 'A', 'selected': True, 'marker_text': 'X', 'source_ref': {'line_id': 'line-2'}},
                                {'label': 'B', 'selected': True, 'marker_text': 'X', 'source_ref': {'line_id': 'line-3'}},
                            ],
                            'source_ref': {'line_id': 'line-1'},
                        }
                    ],
                    'layout_lines': [
                        {'text': 'Question', 'line_id': 'line-1'},
                        {'text': 'X A', 'line_id': 'line-2'},
                        {'text': 'X B', 'line_id': 'line-3'},
                    ],
                    'layout_meta': {},
                    'resolved_fields': [],
                }
            ),
            warnings=[],
        )

        explainability = _build_form_explainability(parsed)

        self.assertIsNotNone(explainability)
        self.assertTrue(explainability['repair_plan']['recommended'])
        self.assertEqual(explainability['repair_plan']['trigger_stage'], 'generic_form_understanding')
        self.assertTrue(any(action['kind'] == 'review_group_selection' for action in explainability['repair_plan']['actions']))

    def test_form_explainability_flags_low_confidence_group(self) -> None:
        parsed = ParsedFile(
            file_name='generic_low_confidence.txt',
            file_type='txt',
            columns=[],
            rows=[],
            content_type='form',
            document_mode='form_layout_mode',
            extraction_status='text_extracted',
            raw_text='Question\nOption A\n',
            text_blocks=[],
            sections=[],
            kv_pairs=[],
            source_candidates=[],
            sheets=[],
            form_model=_build_form_model(
                {
                    'scalars': [],
                    'groups': [
                        {
                            'group_id': 'group_1',
                            'question': 'Question',
                            'group_type': 'single_choice',
                            'group_confidence': 0.54,
                            'selection_confidence': 0.41,
                            'is_ambiguous': False,
                            'options': [
                                {
                                    'label': 'Option A',
                                    'selected': True,
                                    'marker_text': '',
                                    'selection_confidence': 0.41,
                                    'source_ref': {'line_id': 'line-2'},
                                }
                            ],
                            'source_ref': {'line_id': 'line-1'},
                        }
                    ],
                    'layout_lines': [
                        {'text': 'Question', 'line_id': 'line-1'},
                        {'text': 'Option A', 'line_id': 'line-2'},
                    ],
                    'layout_meta': {},
                    'resolved_fields': [],
                }
            ),
            warnings=[],
        )

        explainability = _build_form_explainability(parsed)

        self.assertIsNotNone(explainability)
        self.assertIn('low_confidence_form_groups', explainability['repair_plan']['red_flag_codes'])
        self.assertIn('group_1', explainability['quality_summary']['low_confidence_groups'])
        action = next(action for action in explainability['repair_plan']['actions'] if action.get('group_id') == 'group_1')
        self.assertEqual(action['kind'], 'review_group_selection')
        self.assertEqual(action['priority'], 'medium')
        self.assertAlmostEqual(action['group_confidence'], 0.54)
        self.assertAlmostEqual(action['selection_confidence'], 0.41)

    def test_form_explainability_exposes_pdf_zone_routing(self) -> None:
        parsed = ParsedFile(
            file_name='pdf_explainability.pdf',
            file_type='pdf',
            columns=[],
            rows=[],
            content_type='form',
            document_mode='form_layout_mode',
            extraction_status='text_extracted',
            raw_text='Question\nOption A\n',
            text_blocks=[],
            sections=[],
            kv_pairs=[],
            source_candidates=[],
            sheets=[],
            form_model=_build_form_model(
                {
                    'scalars': [],
                    'groups': [
                        {
                            'group_id': 'group_1',
                            'question': 'Question',
                            'group_type': 'single_choice',
                            'group_confidence': 0.58,
                            'selection_confidence': 0.53,
                            'options': [{'label': 'Option A', 'selected': True, 'source_ref': {'line_id': 'line-2'}}],
                            'source_ref': {'line_id': 'line-1'},
                        }
                    ],
                    'layout_lines': [
                        {'text': 'Question', 'line_id': 'line-1'},
                        {'text': 'Option A', 'line_id': 'line-2'},
                    ],
                    'layout_meta': {},
                    'resolved_fields': [],
                }
            ),
            pdf_zone_summary={
                'dominant_zone': 'form',
                'counts': {'form': 1, 'text': 0, 'noise': 1, 'table': 0},
                'available': True,
                'has_form_zones': True,
                'has_noise_zones': True,
                'best_form_confidence': 0.58,
                'best_noise_confidence': 0.88,
                'has_confident_form_zone': False,
                'low_confidence_form_zones': True,
                'prefer_table_source': False,
                'parser_outputs': {
                    'table': {'zones': []},
                    'form': {'zones': [{'zone_confidence': 0.58}]},
                    'text': {'zones': []},
                    'noise': {'zones': [{'zone_confidence': 0.88}]},
                },
            },
            warnings=[],
        )

        explainability = _build_form_explainability(parsed)

        self.assertIsNotNone(explainability)
        self.assertEqual(explainability['pdf_zone_summary']['routing']['low_confidence_form_zones'], True)
        self.assertIn('low_confidence_form_zones', explainability['repair_plan']['red_flag_codes'])
        self.assertTrue(any(action['kind'] == 'review_pdf_zone_routing' for action in explainability['repair_plan']['actions']))

    def test_form_explainability_exposes_ocr_zone_routing(self) -> None:
        parsed = ParsedFile(
            file_name='photo_explainability.jpg',
            file_type='jpg',
            columns=[],
            rows=[],
            content_type='form',
            document_mode='form_layout_mode',
            extraction_status='text_extracted',
            raw_text='Question\nOption A\n',
            text_blocks=[],
            sections=[],
            kv_pairs=[],
            source_candidates=[],
            sheets=[],
            form_model=_build_form_model(
                {
                    'scalars': [],
                    'groups': [
                        {
                            'group_id': 'group_1',
                            'question': 'Question',
                            'group_type': 'single_choice',
                            'group_confidence': 0.57,
                            'selection_confidence': 0.49,
                            'options': [{'label': 'Option A', 'selected': True, 'source_ref': {'line_id': 'line-2'}}],
                            'source_ref': {'line_id': 'line-1'},
                        }
                    ],
                    'layout_lines': [
                        {'text': 'Question', 'line_id': 'line-1'},
                        {'text': 'Option A', 'line_id': 'line-2'},
                    ],
                    'layout_meta': {},
                    'resolved_fields': [],
                }
            ),
            ocr_used=True,
            ocr_metadata={
                'zone_summary': {
                    'counts': {'form': 1, 'text': 0, 'noise': 2, 'table': 0},
                    'merge_stats': {'input_line_count': 5, 'selected_line_count': 2, 'dropped_low_confidence_lines': 1, 'dropped_noise_lines': 2},
                    'selected_region_ids': ['region_p1_c0_2'],
                    'parser_outputs': {
                        'form': {'regions': [{'zone_id': 'region_p1_c0_2', 'zone_confidence': 0.57}]},
                        'text': {'regions': []},
                        'noise': {'regions': [{'zone_confidence': 0.91}, {'zone_confidence': 0.86}]},
                    },
                }
            },
            warnings=[],
        )

        explainability = _build_form_explainability(parsed)

        self.assertIsNotNone(explainability)
        self.assertEqual(explainability['ocr_zone_summary']['routing']['low_confidence_form_zones'], True)
        self.assertTrue(explainability['ocr_zone_summary']['routing']['noise_dominates'])
        self.assertIn('ocr_noise_dominates', explainability['repair_plan']['red_flag_codes'])
        self.assertIn('ocr_checkbox_selection_review', explainability['repair_plan']['red_flag_codes'])
        self.assertTrue(any(action['kind'] == 'review_ocr_zone_routing' for action in explainability['repair_plan']['actions']))
        self.assertTrue(any(action['kind'] == 'review_ocr_checkbox_selection' for action in explainability['repair_plan']['actions']))

    def test_repair_preview_returns_targeted_patch_for_tax_group(self) -> None:
        parsed = ParsedFile(
            file_name='repair_preview.txt',
            file_type='txt',
            columns=[],
            rows=[],
            content_type='form',
            document_mode='form_layout_mode',
            extraction_status='text_extracted',
            raw_text='',
            text_blocks=[],
            sections=[],
            kv_pairs=[],
            source_candidates=[],
            sheets=[],
            form_model=_build_form_model(
                {
                    'scalars': [],
                    'groups': [
                        {
                            'group_id': 'tax_residency',
                            'question': 'Является ли выгодоприобретатель налоговым резидентом только в РФ',
                            'group_type': 'single_choice',
                            'options': [
                                {
                                    'label': 'ДА',
                                    'selected': True,
                                    'marker_text': 'X',
                                    'source_ref': {'line_id': 'line-2'},
                                },
                                {
                                    'label': 'Не являюсь налоговым резидентом ни в одном государстве',
                                    'selected': True,
                                    'marker_text': 'X',
                                    'source_ref': {'line_id': 'line-3'},
                                },
                            ],
                            'source_ref': {'line_id': 'line-1'},
                        }
                    ],
                    'layout_lines': [
                        {'text': 'Является ли выгодоприобретатель налоговым резидентом только в РФ', 'line_id': 'line-1'},
                        {'text': 'X ДА', 'line_id': 'line-2'},
                        {'text': 'X Не являюсь налоговым резидентом ни в одном государстве', 'line_id': 'line-3'},
                    ],
                    'layout_meta': {'requested_target_fields': ['isResidentRF', 'isTaxResidencyOnlyRF']},
                    'resolved_fields': [],
                }
            ),
            warnings=[],
        )
        payload = RepairPreviewPayload(
            parsed_file=parsed.model_dump() if hasattr(parsed, 'model_dump') else parsed.dict(),
            action={
                'kind': 'repair_ambiguous_resolution',
                'target_field': 'isResidentRF',
                'chunk_refs': {
                    'group_ids': ['tax_residency'],
                    'line_ids': ['line-1', 'line-2', 'line-3'],
                    'scalar_labels': [],
                },
            },
        )

        with patch(
            'form_layout.suggest_form_field_repair',
            return_value=({'status': 'resolved', 'enum_value': 'NOWHERE', 'confidence': 0.77}, []),
        ):
            response = repair_preview(payload)

        self.assertTrue(response['supported'])
        self.assertEqual(response['preview_status'], 'patch_available')
        self.assertEqual(response['proposed_patch']['isResidentRF'], 'NOWHERE')
        self.assertEqual(response['proposed_patch']['isTaxResidencyOnlyRF'], 'NO')
        self.assertEqual(response['local_chunks']['groups'][0]['group_id'], 'tax_residency')
        self.assertEqual([item['field'] for item in response['proposed_resolutions']], ['isResidentRF', 'isTaxResidencyOnlyRF'])

    def test_repair_preview_returns_inspection_only_for_rebuild_action(self) -> None:
        parsed = ParsedFile(
            file_name='repair_preview_generic.txt',
            file_type='txt',
            columns=[],
            rows=[],
            content_type='form',
            document_mode='form_layout_mode',
            extraction_status='text_extracted',
            raw_text='',
            text_blocks=[],
            sections=[],
            kv_pairs=[],
            source_candidates=[],
            sheets=[],
            form_model=_build_form_model(
                {
                    'scalars': [],
                    'groups': [],
                    'layout_lines': [
                        {'text': 'Question', 'line_id': 'line-1'},
                        {'text': 'X A', 'line_id': 'line-2'},
                    ],
                    'layout_meta': {},
                    'resolved_fields': [],
                }
            ),
            warnings=[],
        )
        payload = RepairPreviewPayload(
            parsed_file=parsed.model_dump() if hasattr(parsed, 'model_dump') else parsed.dict(),
            action={
                'kind': 'rebuild_generic_form_understanding',
                'chunk_refs': {
                    'group_ids': [],
                    'line_ids': ['line-1', 'line-2'],
                    'scalar_labels': [],
                },
            },
        )

        response = repair_preview(payload)

        self.assertTrue(response['supported'])
        self.assertEqual(response['preview_status'], 'inspection_only')
        self.assertEqual(response['proposed_patch'], {})
        self.assertEqual(len(response['local_chunks']['lines']), 2)

    def test_repair_apply_updates_local_truth_without_persistence(self) -> None:
        parsed = ParsedFile(
            file_name='repair_apply.txt',
            file_type='txt',
            columns=[],
            rows=[],
            content_type='form',
            document_mode='form_layout_mode',
            extraction_status='text_extracted',
            raw_text='',
            text_blocks=[],
            sections=[],
            kv_pairs=[],
            source_candidates=[],
            sheets=[],
            form_model=_build_form_model(
                {
                    'scalars': [],
                    'groups': [
                        {
                            'group_id': 'tax_residency',
                            'question': 'Является ли выгодоприобретатель налоговым резидентом только в РФ',
                            'group_type': 'single_choice',
                            'options': [
                                {'label': 'ДА', 'selected': True, 'marker_text': 'X', 'source_ref': {'line_id': 'line-2'}},
                                {'label': 'Не являюсь налоговым резидентом ни в одном государстве', 'selected': True, 'marker_text': 'X', 'source_ref': {'line_id': 'line-3'}},
                            ],
                            'source_ref': {'line_id': 'line-1'},
                        }
                    ],
                    'layout_lines': [
                        {'text': 'Является ли выгодоприобретатель налоговым резидентом только в РФ', 'line_id': 'line-1'},
                        {'text': 'X ДА', 'line_id': 'line-2'},
                        {'text': 'X Не являюсь налоговым резидентом ни в одном государстве', 'line_id': 'line-3'},
                    ],
                    'layout_meta': {'requested_target_fields': ['isResidentRF', 'isTaxResidencyOnlyRF']},
                    'resolved_fields': [],
                }
            ),
            warnings=[],
        )
        payload = RepairApplyPayload(
            parsed_file=parsed.model_dump() if hasattr(parsed, 'model_dump') else parsed.dict(),
            action={
                'kind': 'repair_ambiguous_resolution',
                'target_field': 'isResidentRF',
                'chunk_refs': {'group_ids': ['tax_residency'], 'line_ids': ['line-1', 'line-2', 'line-3'], 'scalar_labels': []},
            },
            approved_patch={'isResidentRF': 'NOWHERE', 'isTaxResidencyOnlyRF': 'NO'},
        )

        response = repair_apply(payload, current_user=None)

        self.assertTrue(response['applied'])
        self.assertFalse(response['persistence']['persisted'])
        self.assertEqual(response['parsed_file']['rows'][0]['isResidentRF'], 'NOWHERE')
        self.assertEqual(response['parsed_file']['rows'][0]['isTaxResidencyOnlyRF'], 'NO')
        self.assertEqual(response['form_explainability']['final_source_mode'], 'repair_apply')
        updated_field = next(item for item in response['updated_resolved_fields'] if item['field'] == 'isResidentRF')
        self.assertEqual(updated_field['resolved_by'], 'repair_apply')

    def test_repair_apply_persists_generation_version_when_generation_id_present(self) -> None:
        parsed = ParsedFile(
            file_name='repair_apply_persist.txt',
            file_type='txt',
            columns=[],
            rows=[],
            content_type='form',
            document_mode='form_layout_mode',
            extraction_status='text_extracted',
            raw_text='',
            text_blocks=[],
            sections=[],
            kv_pairs=[],
            source_candidates=[],
            sheets=[],
            form_model=_build_form_model(
                {
                    'scalars': [],
                    'groups': [],
                    'layout_lines': [],
                    'layout_meta': {'requested_target_fields': ['organizationName']},
                    'resolved_fields': [],
                }
            ),
            warnings=[],
        )
        payload = RepairApplyPayload(
            parsed_file=parsed.model_dump() if hasattr(parsed, 'model_dump') else parsed.dict(),
            action={
                'kind': 'repair_scalar_resolution',
                'target_field': 'organizationName',
                'chunk_refs': {'group_ids': [], 'line_ids': [], 'scalar_labels': []},
            },
            approved_patch={'organizationName': 'ООО "Рога и копыта"'},
            generation_id=42,
            notes='apply repair',
        )

        with patch(
            'routes.apply_generation_repair_patch',
            return_value={'generation_id': 42, 'version_id': 9, 'version_number': 3, 'session_id': 17},
        ) as apply_mock:
            response = repair_apply(payload, current_user={'id': 'user-1'})

        self.assertTrue(response['persistence']['persisted'])
        self.assertEqual(response['persistence']['generation_id'], 42)
        self.assertEqual(response['persistence']['version_id'], 9)
        self.assertEqual(response['persistence']['version_number'], 3)
        self.assertEqual(response['persistence']['session_id'], 17)
        apply_mock.assert_called_once()

@unittest.skipIf(openpyxl is None, 'openpyxl is not installed in the current environment')
class ExcelParserTests(unittest.TestCase):
    def test_numeric_excel_headers_are_converted_to_strings(self) -> None:
        fake_workbook = FakeWorkbook(
            {
                'Sheet1': FakeWorksheet(
                    [
                        (1223, 'hsdh', 'sdvsdv'),
                        ('zov', 120, 'sddf'),
                    ]
                )
            }
        )

        with patch('parsers.load_workbook', return_value=fake_workbook):
            path = Path('numeric_headers.xlsx')
            parsed = parse_file(path, path.name)

        self.assertEqual(parsed.columns, ['1223', 'hsdh', 'sdvsdv'])
        self.assertEqual(parsed.rows, [{'1223': 'zov', 'hsdh': 120, 'sdvsdv': 'sddf'}])
        self.assertEqual(len(parsed.sheets), 1)
        self.assertEqual(parsed.sheets[0].name, 'Sheet1')
        self.assertEqual(parsed.sheets[0].columns, ['1223', 'hsdh', 'sdvsdv'])
        self.assertTrue(
            any('Excel first row is treated as column headers' in warning for warning in parsed.warnings)
        )

    def test_multiple_excel_sheets_are_merged(self) -> None:
        fake_workbook = FakeWorkbook(
            {
                'Jan': FakeWorksheet(
                    [
                        ('customerName', 'amount'),
                        ('alice', 10),
                    ]
                ),
                'Feb': FakeWorksheet(
                    [
                        ('customerName', 'amount'),
                        ('bob', 20),
                    ]
                ),
            }
        )

        with patch('parsers.load_workbook', return_value=fake_workbook):
            path = Path('multi_sheet.xlsx')
            parsed = parse_file(path, path.name)

        self.assertEqual(parsed.columns, ['customerName', 'amount'])
        self.assertEqual(
            parsed.rows,
            [
                {'customerName': 'alice', 'amount': 10},
                {'customerName': 'bob', 'amount': 20},
            ],
        )
        self.assertEqual([sheet.name for sheet in parsed.sheets], ['Jan', 'Feb'])
        self.assertEqual(parsed.sheets[0].rows, [{'customerName': 'alice', 'amount': 10}])
        self.assertEqual(parsed.sheets[1].rows, [{'customerName': 'bob', 'amount': 20}])
        self.assertIn('Merged 2 sheets: Jan, Feb', parsed.warnings)

    def test_multiple_tables_in_single_excel_sheet_are_split(self) -> None:
        fake_workbook = FakeWorkbook(
            {
                'Two Tables': FakeWorksheet(
                    [
                        ('TABLE 1: Sales Data', None, None, None, None),
                        ('Product ID', 'Product Name', 'Price', 'Quantity', 'Revenue'),
                        ('P001', 'Laptop', 45000, 5, 225000),
                        ('P002', 'Mouse', 800, 20, 16000),
                        (None, None, None, None, None),
                        ('TABLE 2: Employee Data', None, None, None, None),
                        ('Employee ID', 'Name', 'Department', 'Salary', 'Years Experience'),
                        ('E001', 'John Smith', 'IT', 75000, 5),
                        ('E002', 'Jane Doe', 'Sales', 65000, 3),
                    ]
                )
            }
        )

        with patch('parsers.load_workbook', return_value=fake_workbook):
            path = Path('two_tables.xlsx')
            parsed = parse_file(path, path.name)

        self.assertEqual(
            [sheet.name for sheet in parsed.sheets],
            ['Two Tables · TABLE 1: Sales Data', 'Two Tables · TABLE 2: Employee Data'],
        )
        self.assertEqual(parsed.sheets[0].columns, ['Product ID', 'Product Name', 'Price', 'Quantity', 'Revenue'])
        self.assertEqual(parsed.sheets[1].columns, ['Employee ID', 'Name', 'Department', 'Salary', 'Years Experience'])
        self.assertEqual(parsed.sheets[0].rows[0]['Product Name'], 'Laptop')
        self.assertEqual(parsed.sheets[1].rows[0]['Name'], 'John Smith')
        self.assertTrue(any('Sheet "Two Tables" was split into 2 tables.' in warning for warning in parsed.warnings))

    def test_resolve_draft_json_source_merges_all_excel_sheets_without_selected_sheet(self) -> None:
        parsed = ParsedFile(
            file_name='multi_sheet.xlsx',
            file_type='xlsx',
            columns=['customerName', 'amount', 'status'],
            rows=[
                {'customerName': 'alice', 'amount': 10},
                {'status': 'paid'},
            ],
            sheets=[
                ParsedSheet(name='Jan', columns=['customerName', 'amount'], rows=[{'customerName': 'alice', 'amount': 10}]),
                ParsedSheet(name='Feb', columns=['status'], rows=[{'status': 'paid'}]),
            ],
            warnings=[],
        )

        columns, rows, warnings = resolve_draft_json_source(parsed)

        self.assertEqual(columns, ['customerName', 'amount', 'status'])
        self.assertEqual(rows, [{'customerName': 'alice', 'amount': 10}, {'status': 'paid'}])
        self.assertEqual(warnings, ['Draft JSON uses merged source structure from 2 sheets.'])

    def test_resolve_draft_json_source_uses_selected_sheet(self) -> None:
        parsed = ParsedFile(
            file_name='multi_sheet.xlsx',
            file_type='xlsx',
            columns=['customerName', 'amount', 'status'],
            rows=[
                {'customerName': 'alice', 'amount': 10},
                {'status': 'paid'},
            ],
            sheets=[
                ParsedSheet(name='Jan', columns=['customerName', 'amount'], rows=[{'customerName': 'alice', 'amount': 10}]),
                ParsedSheet(name='Feb', columns=['status'], rows=[{'status': 'paid'}]),
            ],
            warnings=[],
        )

        columns, rows, warnings = resolve_draft_json_source(parsed, selected_sheet='Feb')

        self.assertEqual(columns, ['status'])
        self.assertEqual(rows, [{'status': 'paid'}])
        self.assertEqual(warnings, ['Generated mapping from selected sheet: Feb'])

    def test_resolve_generation_source_uses_selected_sheet(self) -> None:
        parsed = ParsedFile(
            file_name='multi_sheet.xlsx',
            file_type='xlsx',
            columns=['1223', 'hsdh', 'sdvsdv', '345435', '234323', '234'],
            rows=[
                {'1223': 'zov', 'hsdh': 120, 'sdvsdv': 'sddf'},
                {'345435': 'avpva', '234323': 'avp', '234': 'byvapavp'},
            ],
            sheets=[
                ParsedSheet(name='Лист1', columns=['1223', 'hsdh', 'sdvsdv'], rows=[{'1223': 'zov', 'hsdh': 120, 'sdvsdv': 'sddf'}]),
                ParsedSheet(name='Лист2', columns=['345435', '234323', '234'], rows=[{'345435': 'avpva', '234323': 'avp', '234': 'byvapavp'}]),
            ],
            warnings=[],
        )

        columns, rows, warnings = resolve_generation_source(parsed, 'Лист2')

        self.assertEqual(columns, ['345435', '234323', '234'])
        self.assertEqual(rows, [{'345435': 'avpva', '234323': 'avp', '234': 'byvapavp'}])
        self.assertEqual(warnings, ['Generated mapping from selected sheet: Лист2'])

    def test_resolve_generation_source_raises_for_missing_sheet(self) -> None:
        parsed = ParsedFile(
            file_name='multi_sheet.xlsx',
            file_type='xlsx',
            columns=['1223', 'hsdh', 'sdvsdv'],
            rows=[{'1223': 'zov', 'hsdh': 120, 'sdvsdv': 'sddf'}],
            sheets=[ParsedSheet(name='Лист1', columns=['1223', 'hsdh', 'sdvsdv'], rows=[{'1223': 'zov', 'hsdh': 120, 'sdvsdv': 'sddf'}])],
            warnings=[],
        )

        with self.assertRaises(ParseError):
            resolve_generation_source(parsed, 'Лист2')


if __name__ == '__main__':
    unittest.main()
