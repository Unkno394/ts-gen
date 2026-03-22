from __future__ import annotations

import re
from pathlib import Path
from typing import Any

try:
    from docx import Document
except ImportError:  # pragma: no cover - optional dependency in local env
    Document = None

from form_business import (
    _resolve_fatca_group as _resolve_fatca_group_impl,
    _resolve_tax_residency_group as _resolve_tax_residency_group_impl,
    resolve_business_form_fields as _resolve_business_form_fields_impl,
    resolve_form_fields as _resolve_form_fields_impl,
)
from model_client import suggest_form_field_repair


CHECK_MARKERS = {'x', 'X', 'v', 'V', '✓', '✔', '☒', '☑', '■', '1'}
UNCHECKED_MARKERS = {'□', '☐', '○', '◯', '0'}
OCR_CHECKED_MARKER_ALIASES = {'I', 'l', '|', '/', '\\'}
OCR_UNCHECKED_MARKER_ALIASES = {'O', 'o'}
LINE_SPLIT_RE = re.compile(r'[\n\r]+')
SPACE_RE = re.compile(r'\s+')
TOKEN_RE = re.compile(r'[a-zA-Zа-яА-Я0-9]+')
KV_RE = re.compile(r'^(?P<label>[^:]{2,120}?)\s*:\s*(?P<value>.+)$')
PDF_MERGE_Y_THRESHOLD = 6.0
PDF_MERGE_X_THRESHOLD = 32.0
PDF_WRAP_Y_THRESHOLD = 14.0
PDF_WRAP_X_THRESHOLD = 96.0
PDF_ROW_CLUSTER_Y_THRESHOLD = 8.0


def extract_layout_layer(
    *,
    file_path: str | Path,
    file_type: str,
    raw_text: str,
    tables: list[dict[str, Any]] | None = None,
    kv_pairs: list[dict[str, Any]] | None = None,
    text_blocks: list[dict[str, Any]] | None = None,
    sections: list[dict[str, Any]] | None = None,
    layout_blocks: list[dict[str, Any]] | None = None,
) -> dict[str, Any]:
    tables = tables or []
    kv_pairs = kv_pairs or []
    text_blocks = text_blocks or []
    sections = sections or []
    layout_blocks = layout_blocks or []

    seed_scalars: list[dict[str, Any]] = [
        {
            'label': str(pair.get('label') or '').strip(),
            'value': pair.get('value'),
            'source_ref': {
                'source_type': 'line',
                'source_text': str(pair.get('source_text') or ''),
            },
            'confidence': str(pair.get('confidence') or 'high'),
        }
        for pair in kv_pairs
        if str(pair.get('label') or '').strip() and pair.get('value') not in (None, '')
    ]
    layout_lines: list[dict[str, Any]] = []
    raw_table_rows: list[dict[str, Any]] = []
    layout_meta: dict[str, Any] = {
        'file_type': file_type,
        'table_count': len(tables),
        'section_count': len(sections),
        'text_block_count': len(text_blocks),
        'kv_pair_count': len(kv_pairs),
        'reading_order_mode': 'document_order',
        'pipeline_layers': {
            'layout_extraction': {
                'status': 'completed',
                'file_type': file_type,
            },
            'generic_form_understanding': {'status': 'pending'},
            'business_mapping': {'status': 'pending'},
        },
    }

    path = Path(file_path)
    if file_type == 'docx' and Document is not None and path.exists():
        doc = Document(str(path))
        for paragraph_index, paragraph in enumerate(doc.paragraphs):
            text = _clean_text(paragraph.text)
            if not text:
                continue
            layout_lines.append(
                {
                    'text': text,
                    'paragraph_idx': paragraph_index,
                    'source_type': 'paragraph',
                    'tokens': [],
                }
            )
            kv_scalar = _parse_kv_line(text, source_ref={'source_type': 'paragraph', 'paragraph_idx': paragraph_index})
            if kv_scalar is not None:
                seed_scalars.append(kv_scalar)

        table_rows, table_lines, table_meta = _extract_docx_layout_rows(doc)
        raw_table_rows.extend(table_rows)
        layout_lines.extend(table_lines)
        layout_meta.update(table_meta)
        layout_meta['reading_order_mode'] = 'docx_paragraph_table_order'
    else:
        layout_lines.extend(_build_text_layout_lines(raw_text=raw_text, text_blocks=text_blocks, layout_blocks=layout_blocks))
        has_pdf_layout_rows = False
        if file_type == 'pdf':
            pdf_layout_rows = _extract_pdf_layout_rows(layout_lines)
            raw_table_rows.extend(pdf_layout_rows)
            layout_meta['pdf_layout_row_count'] = len(pdf_layout_rows)
            if pdf_layout_rows:
                has_pdf_layout_rows = True
                layout_meta['reading_order_mode'] = 'page_row_cell_order'
        if not has_pdf_layout_rows and any(line.get('page') is not None for line in layout_lines):
            layout_meta['reading_order_mode'] = 'page_column_yx'
        elif layout_lines:
            layout_meta['reading_order_mode'] = 'line_order'

    layout_meta['pipeline_layers']['layout_extraction'].update(
        {
            'layout_line_count': len(layout_lines),
            'raw_table_row_count': len(raw_table_rows),
            'seed_scalar_count': len(seed_scalars),
            'reading_order_mode': layout_meta['reading_order_mode'],
        }
    )
    return {
        'layout_lines': layout_lines,
        'seed_scalars': seed_scalars,
        'raw_table_rows': raw_table_rows,
        'layout_meta': layout_meta,
        'sections': list(sections),
    }


def understand_generic_form(
    *,
    layout_layer: dict[str, Any],
    tables: list[dict[str, Any]] | None = None,
    kv_pairs: list[dict[str, Any]] | None = None,
) -> dict[str, Any] | None:
    tables = tables or []
    kv_pairs = kv_pairs or []
    scalars = [dict(item) for item in layout_layer.get('seed_scalars', []) if isinstance(item, dict)]
    layout_lines = [dict(item) for item in layout_layer.get('layout_lines', []) if isinstance(item, dict)]
    raw_table_rows = [dict(item) for item in layout_layer.get('raw_table_rows', []) if isinstance(item, dict)]
    sections = [dict(item) for item in layout_layer.get('sections', []) if isinstance(item, dict)]
    groups: list[dict[str, Any]] = []
    layout_meta = dict(layout_layer.get('layout_meta') or {})

    if raw_table_rows:
        if _is_pdf_layout_rows(raw_table_rows):
            table_scalars, table_groups = _extract_pdf_form_understanding(raw_table_rows)
        else:
            table_scalars, table_groups = _extract_docx_form_understanding(raw_table_rows)
        scalars.extend(table_scalars)
        groups.extend(table_groups)
        if _has_positioned_layout_lines(layout_lines):
            groups.extend(_extract_groups_from_layout_lines(layout_lines))
            scalars.extend(_extract_scalars_from_layout_lines(layout_lines))
    else:
        groups.extend(_extract_groups_from_layout_lines(layout_lines))
        scalars.extend(_extract_scalars_from_layout_lines(layout_lines))

    groups = _dedupe_groups(groups)
    scalars = _dedupe_scalars(scalars)
    scalars = _filter_noisy_scalars(scalars)
    scalars = _suppress_scalars_consumed_by_groups(scalars, groups)
    groups = _annotate_group_confidence(groups)
    scalars = _annotate_scalar_confidence(scalars)
    section_hierarchy = _build_section_hierarchy(sections, groups)
    structure = _build_form_structure_summary(layout_lines=layout_lines, scalars=scalars, groups=groups, sections=section_hierarchy)

    document_mode = 'data_table_mode'
    if groups or _looks_like_form_layout(kv_pairs=kv_pairs, groups=groups, tables=tables, layout_lines=layout_lines):
        document_mode = 'form_layout_mode'

    if not scalars and not groups and not layout_lines:
        return None

    pipeline_layers = dict(layout_meta.get('pipeline_layers') or {})
    generic_layer = dict(pipeline_layers.get('generic_form_understanding') or {})
    generic_layer.update(
        {
            'status': 'completed',
            'document_mode': document_mode,
            'scalar_count': len(scalars),
            'group_count': len(groups),
            'section_count': len(section_hierarchy),
            'ambiguous_group_count': int(structure.get('ambiguous_group_count') or 0),
        }
    )
    pipeline_layers['generic_form_understanding'] = generic_layer
    layout_meta['pipeline_layers'] = pipeline_layers

    return {
        'scalars': scalars,
        'groups': groups,
        'structure': structure,
        'section_hierarchy': section_hierarchy,
        'layout_lines': layout_lines[:120],
        'layout_meta': {
            **layout_meta,
            'document_mode': document_mode,
        },
        'resolved_fields': [],
    }


def build_form_document_model(
    *,
    file_path: str | Path,
    file_type: str,
    raw_text: str,
    tables: list[dict[str, Any]] | None = None,
    kv_pairs: list[dict[str, Any]] | None = None,
    text_blocks: list[dict[str, Any]] | None = None,
    sections: list[dict[str, Any]] | None = None,
    layout_blocks: list[dict[str, Any]] | None = None,
) -> dict[str, Any] | None:
    layout_layer = extract_layout_layer(
        file_path=file_path,
        file_type=file_type,
        raw_text=raw_text,
        tables=tables,
        kv_pairs=kv_pairs,
        text_blocks=text_blocks,
        sections=sections,
        layout_blocks=layout_blocks,
    )
    return understand_generic_form(
        layout_layer=layout_layer,
        tables=tables,
        kv_pairs=kv_pairs,
    )


def resolve_business_form_fields(
    *,
    form_model: dict[str, Any],
    target_fields: list[Any],
) -> list[dict[str, Any]]:
    return _resolve_business_form_fields_impl(
        form_model=form_model,
        target_fields=target_fields,
        repair_fn=suggest_form_field_repair,
    )


def resolve_form_fields(
    *,
    form_model: dict[str, Any],
    target_fields: list[Any],
) -> list[dict[str, Any]]:
    return _resolve_form_fields_impl(
        form_model=form_model,
        target_fields=target_fields,
        repair_fn=suggest_form_field_repair,
    )


def _extract_docx_layout_rows(doc) -> tuple[list[dict[str, Any]], list[dict[str, Any]], dict[str, Any]]:
    raw_table_rows: list[dict[str, Any]] = []
    layout_lines: list[dict[str, Any]] = []
    form_like_tables = 0

    for table_index, table in enumerate(doc.tables):
        raw_rows: list[dict[str, Any]] = []
        for row_index, row in enumerate(table.rows):
            cell_texts = [_clean_text(cell.text) for cell in row.cells]
            cell_paragraphs: list[list[str]] = []
            for cell_index, cell in enumerate(row.cells):
                paragraph_texts = [_clean_text(paragraph.text) for paragraph in cell.paragraphs if _clean_text(paragraph.text)]
                cell_paragraphs.append(paragraph_texts)
                if paragraph_texts:
                    for paragraph_index, paragraph_text in enumerate(paragraph_texts):
                        layout_lines.append(
                            {
                                'text': paragraph_text,
                                'table_idx': table_index,
                                'row_idx': row_index,
                                'cell_idx': cell_index,
                                'paragraph_idx': paragraph_index,
                                'source_type': 'table_cell',
                                'tokens': [],
                            }
                        )
                    continue
                if not cell_texts[cell_index]:
                    continue
                layout_lines.append(
                    {
                        'text': cell_texts[cell_index],
                        'table_idx': table_index,
                        'row_idx': row_index,
                        'cell_idx': cell_index,
                        'source_type': 'table_cell',
                        'tokens': [],
                    }
                )
            raw_rows.append(
                {
                    'table_idx': table_index,
                    'row_idx': row_index,
                    'cells': cell_texts,
                    'cell_paragraphs': cell_paragraphs,
                }
            )

        if _is_form_like_docx_table([list(item.get('cells', [])) for item in raw_rows]):
            form_like_tables += 1
            raw_table_rows.extend(raw_rows)

    return raw_table_rows, layout_lines, {'form_like_table_count': form_like_tables}


def _extract_pdf_layout_rows(layout_lines: list[dict[str, Any]]) -> list[dict[str, Any]]:
    positioned_lines = [
        dict(line)
        for line in layout_lines
        if _clean_text(line.get('text'))
        and line.get('page') is not None
        and _safe_float(line.get('y')) is not None
    ]
    if not positioned_lines:
        return []

    ordered = sorted(
        positioned_lines,
        key=lambda item: (
            int(item.get('page') or 0),
            float(item.get('y') or 0.0),
            float(item.get('x') or 0.0),
        ),
    )

    clusters: list[dict[str, Any]] = []
    for line in ordered:
        page = int(line.get('page') or 0)
        y = float(line.get('y') or 0.0)
        target_cluster = None
        for cluster in reversed(clusters):
            if cluster['page'] != page:
                break
            if abs(cluster['y'] - y) <= PDF_ROW_CLUSTER_Y_THRESHOLD:
                target_cluster = cluster
                break
            if y - cluster['y'] > PDF_ROW_CLUSTER_Y_THRESHOLD:
                break
        if target_cluster is None:
            target_cluster = {'page': page, 'y': y, 'lines': []}
            clusters.append(target_cluster)
        target_cluster['lines'].append(line)
        target_cluster['y'] = min(target_cluster['y'], y)

    raw_rows: list[dict[str, Any]] = []
    page_row_counters: dict[int, int] = {}
    for cluster in clusters:
        page = int(cluster['page'])
        row_idx = page_row_counters.get(page, 0)
        page_row_counters[page] = row_idx + 1
        row_lines = sorted(cluster['lines'], key=lambda item: float(item.get('x') or 0.0))
        cells: list[str] = []
        cell_paragraphs: list[list[str]] = []
        for line in row_lines:
            text = _clean_text(line.get('text'))
            if not text:
                continue
            cells.append(text)
            cell_paragraphs.append([text])
        if not any(cells):
            continue
        raw_rows.append(
            {
                'table_idx': page - 1,
                'row_idx': row_idx,
                'cells': cells,
                'cell_paragraphs': cell_paragraphs,
                'page': page,
                'source_type': 'pdf_layout_row',
                'row_kind': _classify_pdf_layout_row(cells),
            }
        )

    return raw_rows


def _extract_pdf_form_understanding(raw_table_rows: list[dict[str, Any]]) -> tuple[list[dict[str, Any]], list[dict[str, Any]]]:
    grouped_rows: dict[int, list[dict[str, Any]]] = {}
    for item in raw_table_rows:
        table_index = int(item.get('table_idx') or 0)
        grouped_rows.setdefault(table_index, []).append(item)

    scalars: list[dict[str, Any]] = []
    groups: list[dict[str, Any]] = []
    for table_index, items in grouped_rows.items():
        ordered_rows = [
            dict(item)
            for item in sorted(
                items,
                key=lambda row: (
                    int(row.get('page') or 0),
                    int(row.get('row_idx') or 0),
                ),
            )
        ]
        table_scalars, table_groups = _extract_table_rows_as_form(ordered_rows, table_index=table_index)
        for scalar in table_scalars:
            scalar['source_ref'] = {
                **dict(scalar.get('source_ref', {})),
                'page': _find_pdf_row_page(ordered_rows, scalar.get('source_ref', {})),
            }
        for group in table_groups:
            group['source_ref'] = {
                **dict(group.get('source_ref', {})),
                'page': _find_pdf_row_page(ordered_rows, group.get('source_ref', {})),
            }
            options = []
            for option in group.get('options', []):
                options.append(
                    {
                        **dict(option),
                        'source_ref': {
                            **dict(option.get('source_ref', {})),
                            'page': _find_pdf_row_page(ordered_rows, option.get('source_ref', {})),
                        },
                    }
                )
            group['options'] = options
        scalars.extend(table_scalars)
        groups.extend(table_groups)

    return scalars, groups


def _find_pdf_row_page(raw_rows: list[dict[str, Any]], source_ref: dict[str, Any]) -> int | None:
    row_idx = source_ref.get('row_idx')
    if row_idx is None:
        return None
    for row in raw_rows:
        if row.get('row_idx') == row_idx:
            return int(row.get('page') or 0) or None
    return None


def _classify_pdf_layout_row(cells: list[str]) -> str:
    non_empty = [_clean_text(cell) for cell in cells if _clean_text(cell)]
    if not non_empty:
        return 'empty'

    if len(non_empty) == 2 and _looks_like_heading_fragment_pair(non_empty[0], non_empty[1]):
        return 'header'
    if any(_looks_like_section_heading_text(cell) for cell in non_empty):
        return 'header'

    question_like = any(_looks_like_question_anchor(cell) for cell in non_empty)
    if len(non_empty) == 2 and not question_like and not any(_is_marker_cell(cell) for cell in non_empty):
        return 'field'
    explicit_option_like = any(_extract_option_from_text_line(cell, source_ref={}) is not None for cell in non_empty)
    unmarked_option_like = any(_looks_like_pdf_option_start(cell) for cell in non_empty)

    if question_like and (explicit_option_like or unmarked_option_like):
        return 'mixed'
    if explicit_option_like or unmarked_option_like:
        return 'option'
    if question_like:
        return 'question'
    if len(non_empty) == 2:
        return 'field'
    return 'other'


def _is_pdf_layout_rows(raw_table_rows: list[dict[str, Any]]) -> bool:
    return any(str(item.get('source_type') or '') == 'pdf_layout_row' for item in raw_table_rows)


def _extract_docx_form_understanding(raw_table_rows: list[dict[str, Any]]) -> tuple[list[dict[str, Any]], list[dict[str, Any]]]:
    grouped_rows: dict[int, list[dict[str, Any]]] = {}
    for item in raw_table_rows:
        table_index = int(item.get('table_idx') or 0)
        grouped_rows.setdefault(table_index, []).append(item)

    scalars: list[dict[str, Any]] = []
    groups: list[dict[str, Any]] = []
    for table_index, items in grouped_rows.items():
        ordered_rows = [
            dict(item)
            for item in sorted(items, key=lambda row: int(row.get('row_idx') or 0))
        ]
        table_scalars, table_groups = _extract_table_rows_as_form(ordered_rows, table_index=table_index)
        scalars.extend(table_scalars)
        groups.extend(table_groups)
    return scalars, groups


def _extract_table_rows_as_form(raw_rows: list[dict[str, Any]], *, table_index: int) -> tuple[list[dict[str, Any]], list[dict[str, Any]]]:
    scalars: list[dict[str, Any]] = []
    groups: list[dict[str, Any]] = []
    current_group: dict[str, Any] | None = None

    for row_position, row in enumerate(raw_rows):
        row_index = int(row.get('row_idx') or row_position)
        cells = [_clean_text(cell) for cell in row.get('cells', [])]
        cell_paragraphs = [
            [_clean_text(paragraph) for paragraph in paragraphs if _clean_text(paragraph)]
            for paragraphs in row.get('cell_paragraphs', [])
        ]
        non_empty = [cell for cell in cells if cell]
        if not non_empty:
            continue

        inline_group = _extract_inline_group_from_table_row(
            cells,
            cell_paragraphs=cell_paragraphs,
            table_index=table_index,
            row_index=row_index,
        )
        if inline_group is not None:
            if current_group is not None and current_group.get('options'):
                groups.append(current_group)
            current_group = None
            groups.append(inline_group)
            continue

        scalar_candidate = _extract_scalar_from_table_row(cells, table_index=table_index, row_index=row_index)
        if scalar_candidate is not None:
            if current_group is not None and current_group.get('options'):
                groups.append(current_group)
                current_group = None
            scalars.append(scalar_candidate)
            continue

        option_candidate = _extract_option_from_table_row(cells, table_index=table_index, row_index=row_index)
        if option_candidate is not None:
            if current_group is None:
                current_group = {
                    'group_id': f'group_{table_index}_{row_index}',
                    'question': 'Unknown question',
                    'group_type': 'unknown',
                    'options': [],
                    'source_ref': {'table_idx': table_index, 'row_idx': row_index, 'source_type': 'table_cell'},
                }
            current_group['options'].append(option_candidate)
            current_group['group_type'] = _infer_group_type(current_group)
            continue

        text = ' '.join(non_empty)
        if _looks_like_question_anchor(text):
            if current_group is not None and current_group.get('options'):
                groups.append(current_group)
            current_group = {
                'group_id': _guess_group_id(text, table_index=table_index, row_index=row_index),
                'question': text,
                'group_type': _guess_group_type_from_question(text),
                'options': [],
                'source_ref': {'table_idx': table_index, 'row_idx': row_index, 'source_type': 'table_cell'},
            }
            continue

        if current_group is not None and not current_group.get('options'):
            current_group['question'] = f"{current_group['question']} {text}".strip()

    if current_group is not None and current_group.get('options'):
        groups.append(current_group)

    return scalars, groups


def _extract_inline_group_from_table_row(
    cells: list[str],
    *,
    cell_paragraphs: list[list[str]],
    table_index: int,
    row_index: int,
) -> dict[str, Any] | None:
    non_empty = [(idx, cell) for idx, cell in enumerate(cells) if cell]
    if len(non_empty) < 2:
        return None

    question_idx = next((idx for idx, cell in non_empty if _looks_like_question_anchor(cell)), None)
    if question_idx is None:
        return None

    option_idx = next((idx for idx, _cell in non_empty if idx != question_idx), None)
    if option_idx is None:
        return None

    option_paragraphs = cell_paragraphs[option_idx] if option_idx < len(cell_paragraphs) else []
    option_paragraphs = [paragraph for paragraph in option_paragraphs if paragraph]
    option_text = cells[option_idx]

    if len(option_paragraphs) <= 1 and ';' not in option_text and not _contains_marker(option_text):
        return None

    options = _extract_options_from_docx_cell_paragraphs(
        option_paragraphs or [option_text],
        table_index=table_index,
        row_index=row_index,
        cell_idx=option_idx,
    )
    if not options:
        return None

    question = cells[question_idx]
    group_id = _guess_group_id(question, table_index=table_index, row_index=row_index)
    if any('fatca' in _normalize_phrase(option.get('label')) for option in options):
        group_id = 'fatca_beneficiary'
    group = {
        'group_id': group_id,
        'question': question,
        'group_type': _guess_group_type_from_question(question),
        'options': options,
        'source_ref': {
            'table_idx': table_index,
            'row_idx': row_index,
            'cell_idx': question_idx,
            'source_type': 'table_cell',
        },
    }
    group['group_type'] = _infer_group_type(group)
    return group


def _extract_options_from_docx_cell_paragraphs(
    paragraphs: list[str],
    *,
    table_index: int,
    row_index: int,
    cell_idx: int,
) -> list[dict[str, Any]]:
    options: list[dict[str, Any]] = []
    seen_labels: set[str] = set()

    for paragraph_idx, paragraph in enumerate(paragraphs):
        text = _clean_text(paragraph).rstrip(';').strip()
        if not text:
            continue

        option = _extract_option_from_text_line(
            text,
            source_ref={
                'table_idx': table_index,
                'row_idx': row_index,
                'cell_idx': cell_idx,
                'paragraph_idx': paragraph_idx,
                'source_type': 'table_cell',
            },
        )
        if option is None:
            option = {
                'label': text,
                'selected': False,
                'marker_text': '',
                'source_ref': {
                    'table_idx': table_index,
                    'row_idx': row_index,
                    'cell_idx': cell_idx,
                    'paragraph_idx': paragraph_idx,
                    'source_type': 'table_cell',
                },
            }
        else:
            option['label'] = _clean_text(option.get('label')).rstrip(';').strip()

        normalized_label = _normalize_phrase(option.get('label'))
        if not normalized_label or normalized_label in seen_labels:
            continue
        seen_labels.add(normalized_label)
        options.append(option)

    return options


def _build_text_layout_lines(
    *,
    raw_text: str,
    text_blocks: list[dict[str, Any]],
    layout_blocks: list[dict[str, Any]] | None = None,
) -> list[dict[str, Any]]:
    lines: list[dict[str, Any]] = []
    if layout_blocks:
        for index, block in enumerate(layout_blocks):
            text = _clean_text(block.get('text'))
            if not text:
                continue
            tokens = []
            for token in block.get('tokens', []) or []:
                if not isinstance(token, dict):
                    continue
                tokens.append(
                    {
                        'text': _clean_text(token.get('text')),
                        'x': _safe_float(token.get('x')),
                        'y': _safe_float(token.get('y')),
                        'width': _safe_float(token.get('width')),
                        'height': _safe_float(token.get('height')),
                        'source_type': token.get('source_type') or 'line',
                    }
                )
            lines.append(
                {
                    'text': text,
                    'block_id': str(block.get('id') or f'layout-{index + 1}'),
                    'line_id': str(block.get('id') or f'layout-{index + 1}'),
                    'page': block.get('page'),
                    'column_id': block.get('column_id'),
                    'x': _safe_float(block.get('x')),
                    'y': _safe_float(block.get('y')),
                    'width': _safe_float(block.get('width')),
                    'height': _safe_float(block.get('height')),
                    'source_type': block.get('source_type') or 'line',
                    'tokens': tokens,
                }
            )
        return _merge_pdf_multiline_option_lines(_merge_pdf_marker_lines(lines))
    if text_blocks:
        for index, block in enumerate(text_blocks):
            text = _clean_text(block.get('text'))
            if not text:
                continue
            lines.append(
                {
                    'text': text,
                    'block_id': str(block.get('id') or f'block-{index + 1}'),
                    'column_id': None,
                    'source_type': 'paragraph',
                    'tokens': [],
                }
            )
        return lines

    for index, raw_line in enumerate(LINE_SPLIT_RE.split(raw_text or '')):
        text = _clean_text(raw_line)
        if not text:
            continue
        lines.append(
                {
                    'text': text,
                    'line_id': f'line-{index + 1}',
                    'column_id': None,
                    'source_type': 'line',
                    'tokens': [],
                }
            )
    return lines


def _merge_pdf_marker_lines(layout_lines: list[dict[str, Any]]) -> list[dict[str, Any]]:
    merged: list[dict[str, Any]] = []
    consumed: set[int] = set()
    ordered = sorted(
        layout_lines,
        key=lambda item: (
            int(item.get('page') or 0),
            float(item.get('y') or 0.0),
            float(item.get('x') or 0.0),
        ),
    )
    for index, line in enumerate(ordered):
        if index in consumed:
            continue
        text = _clean_text(line.get('text'))
        if not _is_marker_cell(text):
            merged.append(line)
            continue
        best_candidate_index = None
        best_distance = None
        for candidate_index in range(index + 1, len(ordered)):
            candidate = ordered[candidate_index]
            if candidate_index in consumed:
                continue
            if candidate.get('page') != line.get('page'):
                continue
            candidate_text = _clean_text(candidate.get('text'))
            if not candidate_text or _is_marker_cell(candidate_text):
                continue
            y_distance = abs(float(candidate.get('y') or 0.0) - float(line.get('y') or 0.0))
            x_distance = float(candidate.get('x') or 0.0) - float(line.get('x') or 0.0)
            if y_distance > PDF_MERGE_Y_THRESHOLD or x_distance < 0 or x_distance > PDF_MERGE_X_THRESHOLD * 8:
                if y_distance > PDF_MERGE_Y_THRESHOLD:
                    break
                continue
            total_distance = y_distance + x_distance / 100.0
            if best_distance is None or total_distance < best_distance:
                best_distance = total_distance
                best_candidate_index = candidate_index
        if best_candidate_index is None:
            merged.append(line)
            continue
        consumed.add(best_candidate_index)
        candidate = ordered[best_candidate_index]
        merged_tokens = list(line.get('tokens', [])) + list(candidate.get('tokens', []))
        merged.append(
            {
                **candidate,
                'text': f'{text} {candidate.get("text")}'.strip(),
                'column_id': candidate.get('column_id') or line.get('column_id'),
                'x': min(
                    value
                    for value in [line.get('x'), candidate.get('x')]
                    if value is not None
                ),
                'width': _safe_span_width(line, candidate),
                'tokens': merged_tokens,
            }
        )
    return merged


def _merge_pdf_multiline_option_lines(layout_lines: list[dict[str, Any]]) -> list[dict[str, Any]]:
    merged: list[dict[str, Any]] = []
    ordered = sorted(
        layout_lines,
        key=lambda item: (
            int(item.get('page') or 0),
            int(item.get('column_id') or 0),
            float(item.get('y') or 0.0),
            float(item.get('x') or 0.0),
        ),
    )

    index = 0
    while index < len(ordered):
        current = dict(ordered[index])
        current_text = _clean_text(current.get('text'))
        if not current_text:
            index += 1
            continue

        next_index = index + 1
        while next_index < len(ordered):
            candidate = ordered[next_index]
            if not _should_merge_wrapped_option_line(current, candidate):
                break
            current = {
                **current,
                'text': f'{_clean_text(current.get("text"))} {_clean_text(candidate.get("text"))}'.strip(),
                'width': _safe_span_width(current, candidate),
                'height': max(
                    float(current.get('height') or 0.0),
                    float(candidate.get('height') or 0.0),
                ) or None,
                'tokens': list(current.get('tokens', [])) + list(candidate.get('tokens', [])),
            }
            next_index += 1

        merged.append(current)
        index = next_index

    return merged


def _extract_groups_from_layout_lines(layout_lines: list[dict[str, Any]]) -> list[dict[str, Any]]:
    groups: list[dict[str, Any]] = []
    current_group: dict[str, Any] | None = None

    ordered_lines = sorted(
        layout_lines,
        key=lambda item: (
            int(item.get('page') or 0),
            int(item.get('column_id') or 0),
            float(item.get('y') or 0.0),
            float(item.get('x') or 0.0),
        ),
    )

    for line in ordered_lines:
        text = _clean_text(line.get('text'))
        if not text:
            continue

        option = _extract_contextual_option_from_text_line(text, source_ref=line, current_group=current_group)

        if current_group is not None and _line_breaks_group_scope(current_group, line, option=option):
            if current_group.get('options'):
                groups.append(current_group)
            current_group = None

        if option is not None:
            if current_group is None:
                current_group = {
                    'group_id': _guess_group_id(text),
                    'question': 'Unknown question',
                    'group_type': 'unknown',
                    'options': [],
                    'source_ref': dict(line),
                }
            current_group['options'].append(option)
            current_group['group_type'] = _infer_group_type(current_group)
            continue

        if current_group is not None and current_group.get('options') and _is_group_option_continuation_line(current_group, line):
            last_option = current_group['options'][-1]
            last_option['label'] = f"{_clean_text(last_option.get('label'))} {text}".strip()
            last_option['source_ref'] = {
                **dict(last_option.get('source_ref', {})),
                'continued': True,
            }
            continue

        if _looks_like_question_anchor(text):
            if current_group is not None and current_group.get('options'):
                groups.append(current_group)
            current_group = {
                'group_id': _guess_group_id(text),
                'question': text,
                'group_type': _guess_group_type_from_question(text),
                'options': [],
                'source_ref': dict(line),
            }
            continue

        if current_group is not None and not current_group.get('options'):
            current_group['question'] = f"{current_group['question']} {text}".strip()

    if current_group is not None and current_group.get('options'):
        groups.append(current_group)

    return groups


def _extract_scalars_from_layout_lines(layout_lines: list[dict[str, Any]]) -> list[dict[str, Any]]:
    scalars: list[dict[str, Any]] = []
    table_rows: dict[tuple[int, int], dict[int, str]] = {}
    for line in layout_lines:
        scalar = _parse_kv_line(line.get('text'), source_ref=line)
        if scalar is not None:
            scalars.append(scalar)
        if line.get('source_type') != 'table_cell':
            continue
        table_idx = line.get('table_idx')
        row_idx = line.get('row_idx')
        cell_idx = line.get('cell_idx')
        if not isinstance(table_idx, int) or not isinstance(row_idx, int) or not isinstance(cell_idx, int):
            continue
        text = _clean_text(line.get('text'))
        if not text:
            continue
        table_rows.setdefault((table_idx, row_idx), {})[cell_idx] = text

    for (table_idx, row_idx), row_cells in sorted(table_rows.items()):
        if not row_cells:
            continue
        max_index = max(row_cells)
        cells = [row_cells.get(index, '') for index in range(max_index + 1)]
        scalar = _extract_scalar_from_table_row(cells, table_index=table_idx, row_index=row_idx)
        if scalar is not None:
            scalars.append(scalar)
    return scalars


def _parse_kv_line(text: Any, *, source_ref: dict[str, Any]) -> dict[str, Any] | None:
    cleaned = _clean_text(text)
    if not cleaned:
        return None
    match = KV_RE.match(cleaned)
    if not match:
        return None
    label = _clean_text(match.group('label'))
    value = _clean_text(match.group('value'))
    if not label or not value:
        return None
    if not _is_valid_layout_kv_pair(label, value):
        return None
    return {
        'label': label,
        'value': value,
        'source_ref': dict(source_ref),
        'confidence': 'high',
    }


def _extract_scalar_from_table_row(cells: list[str], *, table_index: int, row_index: int) -> dict[str, Any] | None:
    non_empty = [(idx, cell) for idx, cell in enumerate(cells) if cell]
    if len(non_empty) != 2:
        return None
    if any(_is_marker_cell(cell) for _, cell in non_empty):
        return None
    label_idx, value_idx = non_empty[0][0], non_empty[1][0]
    label, value = non_empty[0][1], non_empty[1][1]
    if _normalize_phrase(label) == _normalize_phrase(value):
        return None
    if label.strip().endswith('?'):
        return None
    if _looks_like_question_anchor(label):
        return None
    if _extract_option_from_text_line(label, source_ref={'table_idx': table_index, 'row_idx': row_index}) is not None:
        return None
    if _looks_like_heading_fragment_pair(label, value):
        return None
    if _extract_option_from_text_line(value, source_ref={'table_idx': table_index, 'row_idx': row_index}) is not None:
        return None
    if len(_tokenize(label)) > 8 or len(value) > 180:
        return None
    return {
        'label': label,
        'value': value,
        'source_ref': {
            'table_idx': table_index,
            'row_idx': row_index,
            'label_cell_idx': label_idx,
            'value_cell_idx': value_idx,
            'source_type': 'table_cell',
        },
        'confidence': 'medium',
    }


def _extract_option_from_table_row(cells: list[str], *, table_index: int, row_index: int) -> dict[str, Any] | None:
    marker_index = next((index for index, cell in enumerate(cells) if _is_marker_cell(cell)), None)
    if marker_index is None:
        return None
    label_parts = [
        cell
        for index, cell in enumerate(cells)
        if index != marker_index and cell and not _is_marker_cell(cell)
    ]
    if not label_parts:
        return None
    label = ' '.join(label_parts).strip()
    if len(label) < 2:
        return None
    marker_text = cells[marker_index]
    return {
        'label': label,
        'selected': _is_checked_marker(marker_text),
        'marker_text': marker_text,
        'source_ref': {
            'table_idx': table_index,
            'row_idx': row_index,
            'marker_cell_idx': marker_index,
            'source_type': 'table_cell',
        },
    }


def _extract_option_from_text_line(text: str, *, source_ref: dict[str, Any]) -> dict[str, Any] | None:
    marker_match = re.match(r'^\s*([XxVv✓✔☒☑□☐○◯Iil|/\\Oo])\s*(.+)$', text)
    if marker_match:
        marker_text = marker_match.group(1)
        return {
            'label': _clean_text(marker_match.group(2)),
            'selected': _is_checked_marker(marker_text),
            'marker_text': marker_text,
            'source_ref': dict(source_ref),
        }
    bracket_match = re.match(r'^\s*\[([XxVv✓✔Iil|/\\Oo ])\]\s*(.+)$', text)
    if bracket_match:
        marker_text = bracket_match.group(1)
        return {
            'label': _clean_text(bracket_match.group(2)),
            'selected': _is_checked_marker(marker_text),
            'marker_text': marker_text,
            'source_ref': dict(source_ref),
        }
    return None


def _extract_contextual_option_from_text_line(
    text: str,
    *,
    source_ref: dict[str, Any],
    current_group: dict[str, Any] | None,
) -> dict[str, Any] | None:
    explicit = _extract_option_from_text_line(text, source_ref=source_ref)
    if explicit is not None:
        explicit['selection_confidence'] = 0.96 if explicit.get('selected') else 0.88
        return explicit
    if current_group is None:
        return None
    if not _looks_like_unmarked_group_option(text, current_group=current_group, source_ref=source_ref):
        return None
    return {
        'label': _clean_text(text),
        'selected': False,
        'marker_text': '',
        'source_ref': dict(source_ref),
        'selection_confidence': 0.45,
    }


def _looks_like_unmarked_group_option(
    text: str,
    *,
    current_group: dict[str, Any],
    source_ref: dict[str, Any],
) -> bool:
    cleaned = _clean_text(text)
    if not cleaned:
        return False
    if _parse_kv_line(cleaned, source_ref={}) is not None:
        return False
    if _looks_like_section_heading_text(cleaned):
        return False
    tokens = _tokenize(cleaned)
    if len(tokens) < 2 or len(tokens) > 18:
        return False

    source_page = source_ref.get('page')
    group_source = dict(current_group.get('source_ref', {}))
    if source_page is not None and group_source.get('page') is not None and source_page != group_source.get('page'):
        return False

    line_column = source_ref.get('column_id')
    group_column = group_source.get('column_id')
    existing_options = [option for option in current_group.get('options', []) if isinstance(option, dict)]
    option_column = None
    if existing_options:
        option_column = dict(existing_options[-1].get('source_ref', {})).get('column_id')

    if option_column is not None and line_column is not None and line_column != option_column:
        return False
    if line_column is not None and group_column is not None and line_column != group_column:
        return True
    if option_column is None and group_column is not None and line_column is not None and line_column == group_column:
        return _looks_like_generic_option_start(cleaned)

    if _looks_like_question_anchor(cleaned):
        return False

    if _looks_like_multi_statement_question(str(current_group.get('question') or '')):
        return True

    first_token = tokens[0]
    if first_token in {'да', 'нет', 'yes', 'no'}:
        return True
    if '%' in cleaned or any(char.isdigit() for char in cleaned):
        return True
    return bool(existing_options)

def _resolve_tax_residency_group(groups: list[dict[str, Any]], *, layout_lines: list[dict[str, Any]]) -> dict[str, Any]:
    return _resolve_tax_residency_group_impl(groups, layout_lines=layout_lines, repair_fn=suggest_form_field_repair)


def _resolve_fatca_group(field_name: str, groups: list[dict[str, Any]], *, layout_lines: list[dict[str, Any]]) -> dict[str, Any]:
    return _resolve_fatca_group_impl(field_name, groups, layout_lines=layout_lines, repair_fn=suggest_form_field_repair)


def _looks_like_form_layout(
    *,
    kv_pairs: list[dict[str, Any]],
    groups: list[dict[str, Any]],
    tables: list[dict[str, Any]],
    layout_lines: list[dict[str, Any]],
) -> bool:
    if groups:
        return True
    if len(kv_pairs) >= 2:
        return True
    if any(_contains_marker(line.get('text')) for line in layout_lines):
        return True
    if len(tables) <= 2 and layout_lines and sum(1 for line in layout_lines if len(_tokenize(line.get('text'))) <= 6) >= 4:
        return True
    return False


def _has_positioned_layout_lines(layout_lines: list[dict[str, Any]]) -> bool:
    return any(line.get('page') is not None for line in layout_lines)


def _looks_like_cross_column_group_attachment(group: dict[str, Any], line: dict[str, Any]) -> bool:
    text = _clean_text(line.get('text'))
    if not text:
        return False
    if _extract_contextual_option_from_text_line(text, source_ref=line, current_group=group) is not None:
        return True
    if group.get('options') and _is_group_option_continuation_line(group, line):
        return True
    return False


def _should_merge_wrapped_option_line(current: dict[str, Any], candidate: dict[str, Any]) -> bool:
    current_text = _clean_text(current.get('text'))
    candidate_text = _clean_text(candidate.get('text'))
    if not current_text or not candidate_text:
        return False
    if current.get('page') != candidate.get('page'):
        return False

    current_column = current.get('column_id')
    candidate_column = candidate.get('column_id')
    if current_column is not None and candidate_column is not None and current_column != candidate_column:
        return False

    if _extract_option_from_text_line(candidate_text, source_ref=candidate) is not None:
        return False
    if _looks_like_pdf_option_start(candidate_text):
        return False
    if _looks_like_question_anchor(candidate_text) or _parse_kv_line(candidate_text, source_ref=candidate) is not None:
        return False
    if _extract_option_from_text_line(current_text, source_ref=current) is None:
        return False

    current_y = _safe_float(current.get('y'))
    candidate_y = _safe_float(candidate.get('y'))
    if current_y is not None and candidate_y is not None and candidate_y - current_y > PDF_WRAP_Y_THRESHOLD:
        return False

    current_x = _safe_float(current.get('x'))
    candidate_x = _safe_float(candidate.get('x'))
    if current_x is not None and candidate_x is not None and abs(candidate_x - current_x) > PDF_WRAP_X_THRESHOLD:
        return False

    return True


def _line_breaks_group_scope(
    group: dict[str, Any],
    line: dict[str, Any],
    *,
    option: dict[str, Any] | None = None,
) -> bool:
    source_ref = dict(group.get('source_ref', {}))
    source_page = source_ref.get('page')
    line_page = line.get('page')
    if source_page is not None and line_page is not None and source_page != line_page:
        return True

    source_column = source_ref.get('column_id')
    line_column = line.get('column_id')
    if source_column is not None and line_column is not None and source_column != line_column:
        if option is not None or _looks_like_cross_column_group_attachment(group, line):
            return False
        return True

    return False


def _is_option_continuation_line(group: dict[str, Any], line: dict[str, Any]) -> bool:
    if not group.get('options'):
        return False
    text = _clean_text(line.get('text'))
    if not text:
        return False
    if _looks_like_question_anchor(text):
        return False
    if _parse_kv_line(text, source_ref=line) is not None:
        return False

    last_option = group['options'][-1]
    last_source = dict(last_option.get('source_ref', {}))
    if last_source.get('page') is not None and line.get('page') not in {None, last_source.get('page')}:
        return False
    if last_source.get('column_id') is not None and line.get('column_id') not in {None, last_source.get('column_id')}:
        return False

    last_y = _safe_float(last_source.get('y'))
    line_y = _safe_float(line.get('y'))
    if last_y is not None and line_y is not None and abs(line_y - last_y) > PDF_WRAP_Y_THRESHOLD:
        return False

    last_x = _safe_float(last_source.get('x'))
    line_x = _safe_float(line.get('x'))
    if last_x is not None and line_x is not None and abs(line_x - last_x) > PDF_WRAP_X_THRESHOLD:
        return False

    return True


def _is_group_option_continuation_line(group: dict[str, Any], line: dict[str, Any]) -> bool:
    if _is_option_continuation_line(group, line):
        text = _clean_text(line.get('text'))
        if not text:
            return True
        last_option = group.get('options', [])[-1]
        last_marker = str(last_option.get('marker_text') or '').strip()
        if last_marker or bool(last_option.get('selected')):
            return False
        if _looks_like_generic_option_start(text):
            return False
        return True

    text = _clean_text(line.get('text'))
    if not text or not group.get('options'):
        return False
    if _looks_like_multi_statement_question(str(group.get('question') or '')) and _looks_like_generic_continuation_fragment(text):
        return True

    if _looks_like_question_anchor(text):
        return False
    if _parse_kv_line(text, source_ref=line) is not None:
        return False
    if _extract_contextual_option_from_text_line(text, source_ref=line, current_group=group) is not None:
        return False

    return False


def _is_form_like_docx_table(raw_rows: list[list[str]]) -> bool:
    if not raw_rows:
        return False
    marker_count = sum(1 for row in raw_rows for cell in row if _is_marker_cell(cell))
    empty_count = sum(1 for row in raw_rows for cell in row if not _clean_text(cell))
    total_count = sum(len(row) for row in raw_rows)
    short_text_rows = sum(1 for row in raw_rows if 0 < len([cell for cell in row if _clean_text(cell)]) <= 3)
    if marker_count > 0:
        return True
    if total_count and empty_count / total_count >= 0.3 and short_text_rows >= 2:
        return True
    return False


def _looks_like_question_anchor(text: str) -> bool:
    normalized = _clean_text(text)
    if _looks_like_generic_option_start(normalized):
        return False
    if '?' in normalized and len(_tokenize(normalized)) >= 3:
        return True
    return len(_tokenize(text)) >= 5 and len(text) >= 24


def _guess_group_id(text: str, *, table_index: int | None = None, row_index: int | None = None) -> str:
    normalized = _normalize_phrase(text)
    if 'fatca' in normalized:
        return 'fatca_beneficiary'
    if 'налогов' in normalized and 'резидент' in normalized:
        return 'tax_residency'
    suffix = 'group'
    if table_index is not None and row_index is not None:
        suffix = f'group_{table_index}_{row_index}'
    return suffix


def _guess_group_type_from_question(text: str) -> str:
    normalized = _normalize_phrase(text)
    if _looks_like_multi_statement_question(text):
        return 'multi_choice'
    if any(word in normalized for word in ('один', 'single', 'только')):
        return 'single_choice'
    return 'unknown'


def _infer_group_type(group: dict[str, Any]) -> str:
    guessed = _guess_group_type_from_question(str(group.get('question') or ''))
    if guessed != 'unknown':
        return guessed
    selected_count = sum(1 for option in group.get('options', []) if option.get('selected'))
    if selected_count > 1:
        return 'multi_choice'
    return 'single_choice'


def _looks_like_multi_statement_question(text: Any) -> bool:
    normalized = _normalize_phrase(text)
    return any(
        phrase in normalized
        for phrase in (
            'хотя бы одно',
            'одно из следующих',
            'следующих утверждений',
            'one of the following',
            'following statements',
        )
    )


def _looks_like_generic_option_start(text: str) -> bool:
    cleaned = _clean_text(text)
    if not cleaned:
        return False
    tokens = _tokenize(cleaned)
    if not tokens:
        return False
    if len(tokens) >= 2 and tokens[1] in {'ли'}:
        return False
    if cleaned.endswith('?'):
        return False
    first = tokens[0]
    if first in {'да', 'нет', 'yes', 'no', 'не'}:
        return True
    if '%' in cleaned or any(char.isdigit() for char in cleaned):
        return True
    return False


def _looks_like_generic_continuation_fragment(text: str) -> bool:
    cleaned = _clean_text(text)
    if not cleaned:
        return False
    if _looks_like_generic_option_start(cleaned):
        return False
    tokens = _tokenize(cleaned)
    if len(tokens) > 8:
        return False
    return cleaned[:1].isupper() or cleaned.startswith(('(', '/', '-'))


def _build_section_hierarchy(
    sections: list[dict[str, Any]],
    groups: list[dict[str, Any]],
) -> list[dict[str, Any]]:
    hierarchy: list[dict[str, Any]] = []
    for index, section in enumerate(sections, start=1):
        title = _clean_text(section.get('title') or f'Section {index}')
        text = _clean_text(section.get('text'))
        if not title and not text:
            continue
        normalized_text = _normalize_phrase(text)
        related_groups = [
            str(group.get('group_id') or '')
            for group in groups
            if _normalize_phrase(str(group.get('question') or '')) and _normalize_phrase(str(group.get('question') or '')) in normalized_text
        ]
        hierarchy.append(
            {
                'section_id': f'section_{index}',
                'title': title,
                'level': 1,
                'related_group_ids': related_groups,
            }
        )
    if not hierarchy and groups:
        hierarchy.append(
            {
                'section_id': 'section_form_root',
                'title': 'Form Body',
                'level': 1,
                'related_group_ids': [str(group.get('group_id') or '') for group in groups],
            }
        )
    return hierarchy


def _annotate_scalar_confidence(scalars: list[dict[str, Any]]) -> list[dict[str, Any]]:
    annotated: list[dict[str, Any]] = []
    for scalar in scalars:
        source_ref = dict(scalar.get('source_ref', {}))
        source_type = str(source_ref.get('source_type') or '')
        confidence_band = str(scalar.get('confidence') or 'medium')
        score = 0.72
        if source_type == 'table_cell':
            score = 0.88
        elif source_type == 'paragraph':
            score = 0.7
        elif source_type == 'line':
            score = 0.74
        if confidence_band == 'high':
            score = max(score, 0.85)
        elif confidence_band == 'low':
            score = min(score, 0.58)
        annotated.append(
            {
                **scalar,
                'confidence_score': round(score, 4),
                'ambiguity_reason': None,
            }
        )
    return annotated


def _annotate_group_confidence(groups: list[dict[str, Any]]) -> list[dict[str, Any]]:
    annotated: list[dict[str, Any]] = []
    for group in groups:
        options = [dict(option) for option in group.get('options', []) if isinstance(option, dict)]
        selected_options = [option for option in options if option.get('selected')]
        explicit_marker_count = sum(1 for option in options if str(option.get('marker_text') or '').strip())
        is_ambiguous = False
        ambiguity_reason = None
        if str(group.get('group_type') or 'unknown') == 'single_choice' and len(selected_options) > 1:
            is_ambiguous = True
            ambiguity_reason = 'multiple_selected_options'
        elif not options:
            is_ambiguous = True
            ambiguity_reason = 'group_without_options'

        group_score = 0.45
        if str(group.get('question') or '').strip():
            group_score += 0.15
        if len(options) >= 2:
            group_score += 0.15
        if explicit_marker_count:
            group_score += 0.15
        if group.get('source_ref'):
            group_score += 0.05
        if is_ambiguous:
            group_score -= 0.15
        group_score = max(0.0, min(group_score, 0.99))

        if len(selected_options) == 1:
            selection_score = 0.95 if str(selected_options[0].get('marker_text') or '').strip() else 0.7
        elif len(selected_options) > 1:
            selection_score = 0.25
        else:
            selection_score = 0.0 if not options else 0.35

        annotated_options: list[dict[str, Any]] = []
        for option in options:
            option_score = option.get('selection_confidence')
            if option_score is None:
                option_score = 0.95 if option.get('selected') and str(option.get('marker_text') or '').strip() else 0.55
                if option.get('selected') and not str(option.get('marker_text') or '').strip():
                    option_score = 0.7
                if not option.get('selected') and not str(option.get('marker_text') or '').strip():
                    option_score = 0.45
            annotated_options.append(
                {
                    **option,
                    'selection_confidence': round(float(option_score), 4),
                    'is_ambiguous': is_ambiguous and option.get('selected', False),
                    'ambiguity_reason': ambiguity_reason if is_ambiguous and option.get('selected', False) else None,
                }
            )

        annotated.append(
            {
                **group,
                'options': annotated_options,
                'group_confidence': round(group_score, 4),
                'selection_confidence': round(selection_score, 4),
                'is_ambiguous': is_ambiguous,
                'ambiguity_reason': ambiguity_reason,
            }
        )
    return annotated


def _build_form_structure_summary(
    *,
    layout_lines: list[dict[str, Any]],
    scalars: list[dict[str, Any]],
    groups: list[dict[str, Any]],
    sections: list[dict[str, Any]],
) -> dict[str, Any]:
    option_count = sum(len(list(group.get('options') or [])) for group in groups)
    selected_option_count = sum(
        1
        for group in groups
        for option in list(group.get('options') or [])
        if isinstance(option, dict) and option.get('selected')
    )
    ambiguous_group_count = sum(1 for group in groups if group.get('is_ambiguous'))
    return {
        'layout_line_count': len(layout_lines),
        'scalar_count': len(scalars),
        'group_count': len(groups),
        'option_count': option_count,
        'selected_option_count': selected_option_count,
        'section_count': len(sections),
        'ambiguous_group_count': ambiguous_group_count,
    }


def _dedupe_scalars(scalars: list[dict[str, Any]]) -> list[dict[str, Any]]:
    result: list[dict[str, Any]] = []
    seen: set[tuple[str, str]] = set()
    for scalar in scalars:
        label = str(scalar.get('label') or '').strip()
        value = str(scalar.get('value') or '').strip()
        key = (_normalize_phrase(label), value.casefold())
        if not label or not value or key in seen:
            continue
        seen.add(key)
        result.append(scalar)
    return result


def _suppress_scalars_consumed_by_groups(
    scalars: list[dict[str, Any]],
    groups: list[dict[str, Any]],
) -> list[dict[str, Any]]:
    if not scalars or not groups:
        return scalars

    consumed_source_refs: list[dict[str, Any]] = []
    consumed_option_texts: set[str] = set()
    consumed_question_texts: set[str] = set()

    for group in groups:
        source_ref = dict(group.get('source_ref', {}))
        if source_ref:
            consumed_source_refs.append(source_ref)

        question_text = _normalize_phrase(group.get('question'))
        if question_text:
            consumed_question_texts.add(question_text)

        for option in group.get('options', []):
            option_ref = dict(option.get('source_ref', {}))
            if option_ref:
                consumed_source_refs.append(option_ref)

            option_text = _normalize_phrase(option.get('label'))
            if option_text:
                consumed_option_texts.add(option_text)

    filtered: list[dict[str, Any]] = []
    for scalar in scalars:
        label = str(scalar.get('label') or '').strip()
        value = str(scalar.get('value') or '').strip()
        normalized_label = _normalize_phrase(label)
        normalized_value = _normalize_phrase(value)
        source_ref = dict(scalar.get('source_ref', {}))

        if any(_source_refs_overlap(source_ref, consumed_ref) for consumed_ref in consumed_source_refs):
            continue
        if normalized_label and normalized_label in consumed_option_texts:
            continue
        if normalized_value and normalized_value in consumed_option_texts:
            continue
        if normalized_label and normalized_label in consumed_question_texts:
            continue
        if normalized_value and normalized_value in consumed_question_texts:
            continue

        filtered.append(scalar)

    return filtered


def _source_refs_overlap(left: dict[str, Any], right: dict[str, Any]) -> bool:
    if not left or not right:
        return False

    shared_scalar_keys = ('line_id', 'block_id', 'paragraph_idx', 'cell_idx')
    for key in shared_scalar_keys:
        left_value = left.get(key)
        right_value = right.get(key)
        if left_value is not None and right_value is not None and left_value == right_value:
            if _shared_parent_context_matches(left, right):
                return True

    table_left = left.get('table_idx')
    table_right = right.get('table_idx')
    row_left = left.get('row_idx')
    row_right = right.get('row_idx')
    if table_left is not None and table_right is not None and row_left is not None and row_right is not None:
        return table_left == table_right and row_left == row_right

    page_left = left.get('page')
    page_right = right.get('page')
    line_left = left.get('line_id')
    line_right = right.get('line_id')
    if page_left is not None and page_right is not None and line_left is not None and line_right is not None:
        return page_left == page_right and line_left == line_right

    return False


def _shared_parent_context_matches(left: dict[str, Any], right: dict[str, Any]) -> bool:
    table_left = left.get('table_idx')
    table_right = right.get('table_idx')
    row_left = left.get('row_idx')
    row_right = right.get('row_idx')
    if table_left is not None and table_right is not None and row_left is not None and row_right is not None:
        return table_left == table_right and row_left == row_right

    page_left = left.get('page')
    page_right = right.get('page')
    if page_left is not None and page_right is not None:
        return page_left == page_right

    return True


def _dedupe_groups(groups: list[dict[str, Any]]) -> list[dict[str, Any]]:
    result: list[dict[str, Any]] = []
    seen: set[tuple[str, tuple[str, ...]]] = set()
    for group in groups:
        question = str(group.get('question') or '').strip()
        options = tuple(_normalize_phrase(option.get('label')) for option in group.get('options', []))
        key = (_normalize_phrase(question), options)
        if not question or not options or key in seen:
            continue
        seen.add(key)
        result.append(group)
    return result


def _filter_noisy_scalars(scalars: list[dict[str, Any]]) -> list[dict[str, Any]]:
    filtered: list[dict[str, Any]] = []
    for scalar in scalars:
        label = str(scalar.get('label') or '').strip()
        value = str(scalar.get('value') or '').strip()
        if not _is_useful_scalar_candidate(label, value):
            continue
        filtered.append(scalar)
    return filtered


def _is_useful_scalar_candidate(label: str, value: str) -> bool:
    normalized_label = _clean_text(label)
    normalized_value = _clean_text(value)
    if not normalized_label or not normalized_value:
        return False
    if _looks_like_heading_fragment_pair(normalized_label, normalized_value):
        return False
    return True


def _is_valid_layout_kv_pair(label: str, value: str) -> bool:
    normalized_label = _clean_text(label)
    normalized_value = _clean_text(value)
    if not normalized_label or not normalized_value:
        return False
    if len(normalized_label) > 80:
        return False
    if len(_tokenize(normalized_label)) > 8:
        return False
    if len(normalized_value) > 260:
        return False
    if _looks_like_heading_fragment_pair(normalized_label, normalized_value):
        return False
    return True


def _looks_like_heading_fragment_pair(label: str, value: str) -> bool:
    if not _looks_like_section_heading_text(label):
        return False
    if any(ch.isdigit() for ch in f'{label} {value}'):
        return False
    if len(_tokenize(value)) < 2:
        return False
    return _uppercase_ratio(value) >= 0.55 or _looks_like_section_heading_text(value)


def _looks_like_section_heading_text(text: str) -> bool:
    tokens = _tokenize(text)
    if len(tokens) < 2:
        return False
    heading_keywords = {
        'сведения',
        'информация',
        'данные',
        'анкета',
        'раздел',
        'часть',
        'приложение',
    }
    if tokens[0] in heading_keywords:
        return True
    return _uppercase_ratio(text) >= 0.9 and len(tokens) >= 4


def _looks_like_pdf_option_start(text: str) -> bool:
    return _looks_like_generic_option_start(text)


def _uppercase_ratio(text: str) -> float:
    letters = [char for char in str(text or '') if char.isalpha()]
    if not letters:
        return 0.0
    uppercase = sum(1 for char in letters if char.isupper())
    return uppercase / len(letters)


def _clean_text(value: Any) -> str:
    return SPACE_RE.sub(' ', str(value or '').replace('\xa0', ' ')).strip()


def _normalize_phrase(value: Any) -> str:
    return ' '.join(_tokenize(value))


def _tokenize(value: Any) -> list[str]:
    return [token.lower() for token in TOKEN_RE.findall(str(value or '')) if token]


def _token_overlap_score(left: list[str], right: list[str]) -> float:
    if not left or not right:
        return 0.0
    left_set = set(left)
    right_set = set(right)
    return len(left_set & right_set) / max(len(left_set), len(right_set), 1)


def _token_stem_overlap_score(left: list[str], right: list[str]) -> float:
    if not left or not right:
        return 0.0
    left_stems = {token[:5] for token in left if token}
    right_stems = {token[:5] for token in right if token}
    if not left_stems or not right_stems:
        return 0.0
    return len(left_stems & right_stems) / max(len(left_stems), len(right_stems), 1)


def _phrase_similarity(left: list[str], right: list[str]) -> float:
    if not left or not right:
        return 0.0
    left_text = ' '.join(left)
    right_text = ' '.join(right)
    if left_text == right_text:
        return 1.0
    overlap = _token_overlap_score(left, right)
    stem_overlap = _token_stem_overlap_score(left, right)
    contains_bonus = 0.15 if left_text in right_text or right_text in left_text else 0.0
    return min(max(overlap, stem_overlap) + contains_bonus, 1.0)


def _contains_marker(value: Any) -> bool:
    text = _canonicalize_marker_text(value)
    if not text:
        return False
    if text in CHECK_MARKERS or text in UNCHECKED_MARKERS:
        return True
    return bool(re.search(r'\[[XxVv✓✔Iil|/\\Oo ]\]', text))


def _is_marker_cell(value: Any) -> bool:
    text = _canonicalize_marker_text(value)
    return text in CHECK_MARKERS or text in UNCHECKED_MARKERS


def _is_checked_marker(value: Any) -> bool:
    return _canonicalize_marker_text(value) in CHECK_MARKERS


def _canonicalize_marker_text(value: Any) -> str:
    text = _clean_text(value)
    if not text:
        return ''
    if text in CHECK_MARKERS or text in UNCHECKED_MARKERS:
        return text
    if text in OCR_CHECKED_MARKER_ALIASES:
        return 'X'
    if text in OCR_UNCHECKED_MARKER_ALIASES:
        return '0'
    return text


def _safe_float(value: Any) -> float | None:
    try:
        return float(value)
    except (TypeError, ValueError):
        return None


def _safe_span_width(left: dict[str, Any], right: dict[str, Any]) -> float | None:
    left_x = _safe_float(left.get('x'))
    right_x = _safe_float(right.get('x'))
    left_width = _safe_float(left.get('width'))
    right_width = _safe_float(right.get('width'))
    if left_x is None and right_x is None:
        return None
    min_x = min(value for value in (left_x, right_x) if value is not None)
    max_x = max(
        value
        for value in (
            None if left_x is None or left_width is None else left_x + left_width,
            None if right_x is None or right_width is None else right_x + right_width,
        )
        if value is not None
    )
    return max(max_x - min_x, 0.0)
