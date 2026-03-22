from __future__ import annotations

import re
from pathlib import Path
from typing import Any

from docx import Document


TABLE_MIN_ROWS = 2
_KV_PAIR_RE = re.compile(
    r'(?P<key>[A-Za-zА-Яа-я0-9_()/ .-]{2,80}?)\s*:\s*(?P<value>.+?)(?=(?:,\s*[A-Za-zА-Яа-я0-9_()/ .-]{2,80}?\s*:)|$)'
)
_DATE_ISO_RE = re.compile(r'^\d{4}-\d{2}-\d{2}(?:[T ][0-9:.]+Z?)?$')
_DATE_DOTTED_RE = re.compile(r'^\d{2}\.\d{2}\.\d{4}$')
_NUMBER_RE = re.compile(r'^[+-]?\d+(?:[.,]\d+)?$')
_ORG_RE = re.compile(r'^(?:ООО|АО|ОАО|ЗАО|ПАО|ИП)\b', re.IGNORECASE)


def clean_text(value: Any) -> str:
    if value is None:
        return ""
    return " ".join(str(value).replace("\xa0", " ").split()).strip()


def table_to_rows(table) -> tuple[list[str], list[dict[str, str]]] | None:
    raw_rows: list[list[str]] = []
    for row in table.rows:
        cells = [clean_text(cell.text) for cell in row.cells]
        if any(cells):
            raw_rows.append(cells)

    if len(raw_rows) < TABLE_MIN_ROWS:
        return None

    header = raw_rows[0]
    if not any(header):
        return None

    columns = [col if col else f"column{idx + 1}" for idx, col in enumerate(header)]
    rows: list[dict[str, str]] = []

    for raw in raw_rows[1:]:
        padded = raw + [""] * max(0, len(columns) - len(raw))
        row_obj = {columns[idx]: padded[idx] for idx in range(len(columns))}
        if any(v != "" for v in row_obj.values()):
            rows.append(row_obj)

    return columns, rows


def _extract_kv_rows_from_paragraphs(paragraphs: list[str]) -> tuple[list[str], list[dict[str, str]]] | None:
    candidate_rows: list[dict[str, str]] = []
    column_order: list[str] = []
    seen_columns: set[str] = set()

    for paragraph in paragraphs:
        matches = list(_KV_PAIR_RE.finditer(paragraph))
        if len(matches) < 2:
            continue

        row: dict[str, str] = {}
        for match in matches:
            key = clean_text(match.group('key')).strip(' ,;')
            value = clean_text(match.group('value')).strip(' ,;')
            if not key or not value:
                continue
            row[key] = value
            if key not in seen_columns:
                seen_columns.add(key)
                column_order.append(key)
        if len(row) >= 2:
            candidate_rows.append(row)

    if len(candidate_rows) < 2 or len(column_order) < 2:
        return None

    rows = [{column: row.get(column, '') for column in column_order} for row in candidate_rows]
    return column_order, rows


def _extract_delimited_rows_from_paragraphs(paragraphs: list[str]) -> tuple[list[str], list[dict[str, str]]] | None:
    for paragraph in paragraphs:
        if '|' not in paragraph:
            continue
        cells = [clean_text(part) for part in paragraph.split('|') if clean_text(part)]
        if len(cells) < 4:
            continue

        width = _infer_repeated_row_width(cells)
        if width is None:
            continue

        raw_rows = [cells[index:index + width] for index in range(0, len(cells), width)]
        if len(raw_rows) < 2:
            continue

        columns = _infer_repeated_row_columns(raw_rows)
        rows = [{columns[index]: row[index] for index in range(len(columns))} for row in raw_rows if len(row) == len(columns)]
        if len(rows) >= 2:
            return columns, rows

    return None


def _infer_repeated_row_width(cells: list[str]) -> int | None:
    total = len(cells)
    best_width: int | None = None
    best_score = -1.0

    for width in range(2, min(12, total // 2) + 1):
        if total % width != 0:
            continue
        row_count = total // width
        if row_count < 2:
            continue
        rows = [cells[index:index + width] for index in range(0, total, width)]
        score = _score_repeated_rows(rows)
        if score > best_score:
            best_score = score
            best_width = width

    return best_width


def _score_repeated_rows(rows: list[list[str]]) -> float:
    row_count = len(rows)
    width = len(rows[0]) if rows else 0
    if row_count < 2 or width < 2:
        return -1.0

    score = 0.0
    for column_index in range(width):
        bucket: dict[str, int] = {}
        for row in rows:
            cell_type = _classify_delimited_cell(row[column_index])
            bucket[cell_type] = bucket.get(cell_type, 0) + 1
        dominant = max(bucket.values()) if bucket else 0
        score += dominant / row_count
        if len(bucket) == 1 and 'text' not in bucket:
            score += 0.5
    return score


def _classify_delimited_cell(value: str) -> str:
    text = clean_text(value)
    lowered = text.casefold()
    if _DATE_ISO_RE.fullmatch(text) or _DATE_DOTTED_RE.fullmatch(text):
        return 'date'
    if _NUMBER_RE.fullmatch(text):
        return 'number'
    if lowered.startswith(('deal_', 'deal-', 'sales-')):
        return 'identifier'
    if _ORG_RE.match(text):
        return 'organization'
    parts = text.split()
    if len(parts) >= 2 and all(part and part[0].isupper() for part in parts[:2]):
        return 'person'
    return 'text'


def _infer_repeated_row_columns(rows: list[list[str]]) -> list[str]:
    inferred_types: list[str] = []
    width = len(rows[0]) if rows else 0
    for column_index in range(width):
        bucket: dict[str, int] = {}
        for row in rows:
            cell_type = _classify_delimited_cell(row[column_index])
            bucket[cell_type] = bucket.get(cell_type, 0) + 1
        inferred_types.append(max(bucket, key=bucket.get) if bucket else 'text')

    columns: list[str] = []
    used_names: dict[str, int] = {}
    for index, cell_type in enumerate(inferred_types):
        name = _column_name_from_context(index=index, total=width, cell_type=cell_type, inferred_types=inferred_types)
        used_count = used_names.get(name, 0)
        used_names[name] = used_count + 1
        columns.append(name if used_count == 0 else f'{name}{used_count + 1}')
    return columns


def _column_name_from_context(*, index: int, total: int, cell_type: str, inferred_types: list[str]) -> str:
    if cell_type == 'identifier':
        return 'identifier'
    if cell_type == 'date':
        return 'date'
    if cell_type == 'number':
        return 'amount'
    if cell_type == 'organization':
        return 'organization'
    if cell_type == 'person':
        return 'responsible'
    if total >= 6:
        if index == 1:
            return 'name'
        if index == 3:
            return 'stage'
        if index == total - 2 and inferred_types[-1] == 'person':
            return 'product'
    return f'column{index + 1}'


def parse_docx(path: str | Path) -> dict[str, Any]:
    doc = Document(str(path))

    warnings: list[str] = []
    paragraphs = [clean_text(p.text) for p in doc.paragraphs if clean_text(p.text)]
    tables_data: list[tuple[list[str], list[dict[str, str]]]] = []

    for table in doc.tables:
        parsed = table_to_rows(table)
        if parsed:
            tables_data.append(parsed)

    if tables_data:
        columns, rows = tables_data[0]
        if len(tables_data) > 1:
            warnings.append(f"Found {len(tables_data)} tables in DOCX.")
        return {
            "file_name": Path(path).name,
            "file_type": "docx",
            "content_type": "table",
            "columns": columns,
            "rows": rows,
            "tables": [
                {
                    "name": f"Table {index + 1}",
                    "columns": item_columns,
                    "rows": item_rows,
                }
                for index, (item_columns, item_rows) in enumerate(tables_data)
            ],
            "text": "\n".join(paragraphs),
            "blocks": [{"type": "paragraph", "text": text} for text in paragraphs],
            "warnings": warnings,
        }

    reconstructed = _extract_kv_rows_from_paragraphs(paragraphs)
    if reconstructed is None:
        reconstructed = _extract_delimited_rows_from_paragraphs(paragraphs)

    if reconstructed is not None:
        columns, rows = reconstructed
        warnings.append("No explicit DOCX table found. Reconstructed tabular preview from inline paragraph data.")
        return {
            "file_name": Path(path).name,
            "file_type": "docx",
            "content_type": "table",
            "columns": columns,
            "rows": rows,
            "tables": [
                {
                    "name": "Reconstructed Table 1",
                    "columns": columns,
                    "rows": rows,
                }
            ],
            "text": "\n".join(paragraphs),
            "blocks": [{"type": "paragraph", "text": text} for text in paragraphs],
            "warnings": warnings,
        }

    warnings.append("No tables found in DOCX. Returned text blocks only.")
    return {
        "file_name": Path(path).name,
        "file_type": "docx",
        "content_type": "text",
        "columns": [],
        "rows": [],
        "tables": [],
        "text": "\n".join(paragraphs),
        "blocks": [{"type": "paragraph", "text": text} for text in paragraphs],
        "warnings": warnings,
    }
